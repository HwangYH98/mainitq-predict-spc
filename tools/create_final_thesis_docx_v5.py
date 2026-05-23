from __future__ import annotations

import json
import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor

import create_final_thesis_docx_v4 as v4


ROOT = Path(__file__).resolve().parents[1]
OUTPUTS = ROOT / "outputs"
DOCX_PATH = OUTPUTS / "final_thesis_manuscript_29p_v5.docx"
CHECKLIST_PATH = OUTPUTS / "final_thesis_manuscript_29p_v5_format_checklist.md"
VALIDATION_PATH = OUTPUTS / "final_thesis_manuscript_29p_v5_validation.json"

BODY_FONT = "바탕"
HEADING_FONT = "HY견명조"
FALLBACK_HEADING_FONT = "함초롬바탕"
ACCENT = RGBColor(22, 96, 78)
LIGHT_ACCENT = "DDEFE8"


FIGURE_TOC = [
    "그림 1. 데이터 전처리 및 예측 pipeline",
    "그림 2. XGBoost PR curve",
    "그림 3. Threshold tuning 결과",
    "그림 4. Predictive SPC 위험 확률 관리도",
    "그림 5. MaintiQ Predict 전체 시스템 구조",
    "그림 6. MaintiQ Predict 메인 화면",
    "그림 7. 승인형 작업지시 workflow",
    "그림 8. SHAP summary plot",
    "그림 9. SHAP factor bar chart",
    "그림 10. Operational value simulation",
    "그림 11. SCANIA official cost comparison",
    "그림 12. Public benchmark cost comparison",
    "그림 13. Public benchmark lead-time summary",
]

TABLE_TOC = [
    "표 1. 보전 전략 비교",
    "표 2. FMEA/RPN 개념과 본 시스템 변수의 대응",
    "표 3. 선행연구 흐름과 본 연구 반영점",
    "표 4. AI4I 데이터 전처리 규칙",
    "표 5. AI4I baseline 모델 성능",
    "표 6. SPC-only와 ML+SPC 비교 해석",
    "표 7. GenAI 리포트 생성 근거",
    "표 8. Full/Lite 모드 구분",
    "표 9. 회사 데이터 실증 파일 조합별 claim 상태",
    "표 10. 실험 설계 요약",
    "표 11. SCANIA official cost metric 결과 요약",
    "표 12. 가능한 주장과 금지 주장",
]


def set_font(run, name: str = BODY_FONT, size: int = 11, bold: bool = False, color: RGBColor | None = None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def compact_page_break(doc: Document) -> None:
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


class ThesisV5Builder(v4.ThesisBuilder):
    def _configure(self) -> None:
        super()._configure()
        for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
            style = self.doc.styles[style_name]
            style.font.name = HEADING_FONT
            style._element.rPr.rFonts.set(qn("w:eastAsia"), HEADING_FONT)
            style.paragraph_format.space_before = Pt(6)
            style.paragraph_format.space_after = Pt(2)

    def centered(self, text: str, size: int = 11, bold: bool = False, before: int = 0, after: int = 0, font: str = BODY_FONT) -> None:
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.line_spacing = 1.35
        p.paragraph_format.space_before = Pt(before)
        p.paragraph_format.space_after = Pt(after)
        r = p.add_run(text)
        set_font(r, font, size, bold)

    def title_line(self, text: str, size: int = 17, before: int = 0, after: int = 0) -> None:
        self.centered(text, size=size, bold=True, before=before, after=after, font=HEADING_FONT)

    def simple_form_table(self, rows: list[tuple[str, str]], widths: tuple[float, float] = (3.2, 8.6)) -> None:
        table = self.doc.add_table(rows=0, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"
        for label, value in rows:
            cells = table.add_row().cells
            for idx, cell in enumerate(cells):
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                cell.width = Cm(widths[idx])
                v4.set_cell_margins(cell, 120)
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.first_line_indent = Cm(0)
                r = p.add_run(label if idx == 0 else value)
                set_font(r, BODY_FONT, 11, idx == 0)
        self.doc.add_paragraph()

    def add_formal_cover(self) -> None:
        self.centered("학 사 학 위 논 문", 15, True, before=26, after=32, font=HEADING_FONT)
        self.title_line("AI 예지보전 기반 스마트 제조 설비", 18, before=6)
        self.title_line("위험관리 시스템 구현", 18, after=10)
        self.centered("MaintiQ Predict: ML·SPC·GenAI 통합 접근", 11, False, after=38, font=BODY_FONT)
        self.centered("[제출일]", 12, False, after=30)
        self.centered("전남대학교 공과대학", 13, True, after=6, font=HEADING_FONT)
        self.centered("산업공학과", 13, True, after=6, font=HEADING_FONT)
        self.centered("[성명]", 13, True, after=16, font=HEADING_FONT)
        self.centered("지도교수  [지도교수]", 11, False, after=0)
        compact_page_break(self.doc)

    def add_inner_cover_i(self) -> None:
        self.centered("내표지 I", 11, False, before=0, after=10)
        self.centered("학 사 학 위 논 문", 15, True, before=18, after=28, font=HEADING_FONT)
        self.title_line("AI 예지보전 기반 스마트 제조 설비 위험관리 시스템 구현", 17, after=12)
        self.centered("MaintiQ Predict: ML·SPC·GenAI 통합 접근", 11, False, after=36)
        self.centered("이 논문을 학사학위 논문으로 제출함", 12, False, after=34)
        self.simple_form_table([
            ("소속", "전남대학교 공과대학 산업공학과"),
            ("학번", "[학번]"),
            ("성명", "[성명]"),
            ("지도교수", "[지도교수]"),
            ("제출일", "[제출일]"),
        ])
        compact_page_break(self.doc)

    def add_inner_cover_ii(self) -> None:
        self.centered("내표지 II", 11, False, before=0, after=10)
        self.centered("AI 예지보전 기반 스마트 제조 설비 위험관리 시스템 구현", 15, True, before=20, after=22, font=HEADING_FONT)
        self.centered("위 논문을 학사학위 논문으로 인정함", 12, False, after=22)
        self.simple_form_table([
            ("심사위원장", "[성명]        (인)"),
            ("심사위원", "[성명]        (인)"),
            ("심사위원", "[성명]        (인)"),
            ("지도교수", "[지도교수]    (인)"),
        ])
        self.centered("전남대학교 공과대학 산업공학과", 12, True, before=28, font=HEADING_FONT)
        compact_page_break(self.doc)

    def add_static_list_page(self, title: str, items: list[str]) -> None:
        self.heading1(title)
        for item in items:
            p = self.doc.add_paragraph()
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.line_spacing = 1.6
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(item)
            set_font(r, BODY_FONT, 10)
        compact_page_break(self.doc)

    def add_combined_figure_table_lists(self) -> None:
        self.heading1("그림 목차")
        table = self.doc.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"
        for idx, header in enumerate(["그림 목차", "표 목차"]):
            cell = table.rows[0].cells[idx]
            v4.set_cell_shading(cell, LIGHT_ACCENT)
            v4.set_cell_margins(cell, 90)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Cm(0)
            r = p.add_run(header)
            set_font(r, BODY_FONT, 9, True, ACCENT)
        for row_idx in range(max(len(FIGURE_TOC), len(TABLE_TOC))):
            cells = table.add_row().cells
            for col_idx, items in enumerate([FIGURE_TOC, TABLE_TOC]):
                cell = cells[col_idx]
                v4.set_cell_margins(cell, 65)
                p = cell.paragraphs[0]
                p.paragraph_format.first_line_indent = Cm(0)
                p.paragraph_format.line_spacing = 1.0
                p.paragraph_format.space_after = Pt(0)
                r = p.add_run(items[row_idx] if row_idx < len(items) else "")
                set_font(r, BODY_FONT, 7)
        compact_page_break(self.doc)

    def add_table_of_contents(self) -> None:
        self.heading1("목차")
        items = [
            "국문초록",
            "1. 서론",
            "   가. 연구 배경",
            "   나. 연구 목적과 범위",
            "   다. 연구의 차별성",
            "2. 이론적 배경 및 선행연구",
            "   가. 예지보전과 CBM",
            "   나. SPC와 관리도 이론",
            "   다. FMEA/RPN과 위험 우선순위",
            "   라. 기계학습 기반 고장 예측",
            "   마. SHAP, GenAI, 비용 민감 학습",
            "3. 연구 방법",
            "4. 시스템 구현",
            "5. 실험 및 검증",
            "6. 결론 및 향후 연구",
            "참고문헌",
            "영문초록(Abstract)",
            "부록",
        ]
        for item in items:
            p = self.doc.add_paragraph()
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.line_spacing = 1.45
            p.paragraph_format.space_after = Pt(1)
            r = p.add_run(item)
            set_font(r, BODY_FONT, 10)
        compact_page_break(self.doc)

    def add_korean_abstract(self) -> None:
        self.heading1("국문초록")
        self.paragraph(
            "본 연구는 제조 설비의 고장 위험을 사전에 파악하고 관리자 의사결정으로 연결하기 위한 AI 예지보전 기반 스마트 제조 설비 위험관리 시스템을 구현한다. "
            "기존 연구는 고장 분류 모델 성능, 관리도 기반 이상탐지, 설명가능 인공지능, 정비 의사결정 중 일부 기능을 개별적으로 다루는 경우가 많았다. "
            "본 연구는 AI4I 2020 데이터셋을 기반으로 XGBoost 고장 확률 예측 모델을 학습하고, threshold tuning을 통해 운영 정책에 맞는 위험 판정 기준을 도출하였다. "
            "또한 예측 확률을 시간축에 배치하여 Predictive SPC 관점의 위험 추세를 확인하고, SHAP 요인과 GenAI 관리자 리포트를 결합해 위험 원인과 조치 방향을 요약하였다."
        )
        self.paragraph(
            f"구현 결과물인 MaintiQ Predict는 CSV 입력, 전처리, 고장 확률 산출, 위험 우선순위 표시, AI 리포트 생성, 승인형 작업지시 기록을 하나의 데스크톱 애플리케이션 흐름으로 제공한다. "
            f"AI4I 기준 XGBoost 모델은 PR-AUC {self.evidence['xgb_pr_auc']:.4f}, ROC-AUC {self.evidence['xgb_roc_auc']:.4f}를 보였으며, 선택 threshold {self.evidence['threshold']:.2f}에서 F1-score {self.evidence['tuned_f1']:.4f}를 달성하였다. "
            f"Gemini API 기반 리포트 검증에서는 {self.evidence['genai_mode']} 방식으로 고위험 확률 {self.evidence['genai_prob']:.6f} 사례에 대한 관리자 참고 리포트를 생성하였다. "
            f"추가로 SCANIA Component X 공개 benchmark에서는 official cost metric 기준 rule baseline 대비 약 {self.evidence['scania_cost_improvement'] * 100:.2f}% 개선 가능성을 확인하였다."
        )
        self.paragraph("주요어: 예지보전, 스마트 제조, XGBoost, Predictive SPC, GenAI, 작업지시, Cost-sensitive Learning")
        compact_page_break(self.doc)

    def add_front_matter(self) -> None:
        self.add_formal_cover()
        self.add_inner_cover_i()
        self.add_inner_cover_ii()
        self.add_table_of_contents()
        self.add_combined_figure_table_lists()
        self.add_korean_abstract()

    def add_english_abstract_before_appendix(self) -> None:
        target = None
        for p in self.doc.paragraphs:
            if p.text.strip() == "부록":
                target = p
        if target is None:
            return
        # python-docx does not expose arbitrary insertion before an existing
        # paragraph, so use direct XML construction for this post-body section.
        target._p.addprevious(self._make_paragraph_xml("Abstract", style="Heading1"))
        target._p.addprevious(self._make_paragraph_xml(
            "This study implements MaintiQ Predict, a desktop predictive-maintenance MVP integrating XGBoost prediction, threshold tuning, Predictive SPC, GenAI reporting, and human-approved work orders. AI4I evaluation showed PR-AUC 0.8014, ROC-AUC 0.9736, threshold F1-score 0.7752, and SCANIA Component X showed 17.02% official-cost improvement over a rule baseline; real field deployment remains future work."
        ))

    def compact_appendix(self) -> None:
        target = None
        paragraphs = list(self.doc.paragraphs)
        for p in paragraphs:
            if p.text.strip() == "부록":
                target = p
        if target is None:
            return
        remove_started = False
        for p in list(self.doc.paragraphs):
            if p._p is target._p:
                remove_started = True
                continue
            if remove_started:
                p._element.getparent().remove(p._element)
        self.heading2("가. 제출 전 확인")
        self.paragraph(
            "최종 제출 전 [성명], [학번], [지도교수], [제출일], 표·그림 번호, HWP 양식 서명란을 확인한다. "
            "실제 현장 실증에는 센서값, 고장 라벨, 정비 이력, downtime, 비용 로그가 필요하다."
        )

    def compact_references(self) -> None:
        in_refs = False
        for p in self.doc.paragraphs:
            text = p.text.strip()
            if text == "참고문헌" or text.startswith("Jardine,"):
                in_refs = True
                if text == "참고문헌":
                    continue
            if in_refs and (text == "부록" or text == "Abstract"):
                break
            if in_refs and text:
                p.paragraph_format.first_line_indent = Cm(0)
                p.paragraph_format.line_spacing = 1.0
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    set_font(run, BODY_FONT, 8)
        # Some Windows shells can corrupt non-ASCII literals in ad-hoc runs.
        # This ASCII fallback keeps the compact reference formatting stable.
        ascii_ref_started = False
        for p in self.doc.paragraphs:
            text = p.text.strip()
            if text.startswith("Jardine,"):
                ascii_ref_started = True
            if ascii_ref_started and (text == "Abstract" or text.startswith("This study implements")):
                break
            if ascii_ref_started and text:
                p.paragraph_format.first_line_indent = Cm(0)
                p.paragraph_format.line_spacing = 1.0
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    set_font(run, BODY_FONT, 8)

    def _make_paragraph_xml(self, text: str, style: str | None = None):
        p = OxmlElement("w:p")
        p_pr = OxmlElement("w:pPr")
        if style:
            p_style = OxmlElement("w:pStyle")
            p_style.set(qn("w:val"), style)
            p_pr.append(p_style)
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:line"), "300" if not style else "400")
        spacing.set(qn("w:lineRule"), "auto")
        p_pr.append(spacing)
        p.append(p_pr)
        r = OxmlElement("w:r")
        r_pr = OxmlElement("w:rPr")
        fonts = OxmlElement("w:rFonts")
        fonts.set(qn("w:ascii"), BODY_FONT if not style else HEADING_FONT)
        fonts.set(qn("w:hAnsi"), BODY_FONT if not style else HEADING_FONT)
        fonts.set(qn("w:eastAsia"), BODY_FONT if not style else HEADING_FONT)
        r_pr.append(fonts)
        size = OxmlElement("w:sz")
        size.set(qn("w:val"), "20" if not style else "32")
        r_pr.append(size)
        if style:
            b = OxmlElement("w:b")
            r_pr.append(b)
        r.append(r_pr)
        t = OxmlElement("w:t")
        t.text = text
        r.append(t)
        p.append(r)
        return p

    def _make_page_break_xml(self):
        p = OxmlElement("w:p")
        r = OxmlElement("w:r")
        br = OxmlElement("w:br")
        br.set(qn("w:type"), "page")
        r.append(br)
        p.append(r)
        return p

    def save(self) -> None:
        self.compact_references()
        self.compact_appendix()
        self.add_english_abstract_before_appendix()
        self.doc.save(DOCX_PATH)


def extract_docx_text(path: Path) -> str:
    doc = Document(path)
    parts: list[str] = []
    for p in doc.paragraphs:
        parts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def validate_text(text: str) -> dict:
    forbidden = [
        "실제 PLC/SCADA 배포 완료",
        "실제 회사 데이터 검증 완료",
        "실제 비용 절감 실증 완료",
    ]
    api_patterns = [r"AIza[0-9A-Za-z_-]{20,}", r"sk-[0-9A-Za-z_-]{20,}"]
    required_order = ["학 사 학 위 논 문", "내표지", "목차", "그림 목차", "표 목차", "국문초록", "1. 서론", "참고문헌", "Abstract", "부록"]
    return {
        "char_count_no_space": len(re.sub(r"\s+", "", text)),
        "replacement_character_count": text.count("�"),
        "api_key_pattern_count": sum(len(re.findall(pattern, text)) for pattern in api_patterns),
        "forbidden_phrase_hits": [phrase for phrase in forbidden if phrase in text],
        "required_metric_hits": {m: (m in text) for m in ["0.8014", "0.9736", "0.87", "0.7752", "0.993616", "17.02"]},
        "required_order_hits": {m: (m in text) for m in required_order},
    }


def write_checklist(builder: ThesisV5Builder, validation: dict) -> None:
    lines = [
        "# final_thesis_manuscript_29p_v5 양식 체크리스트",
        "",
        "## 작성법 반영",
        "- A4, 위 35mm / 왼쪽 35mm / 오른쪽 30mm / 아래 25mm / 꼬리말 15mm",
        "- 본문 바탕 11pt, 줄간격 200%",
        "- 제목 HY견명조 지정(시스템에 없으면 Word/HWP에서 대체 표시 가능)",
        "- 순서: 표지 → 내표지 I → 내표지 II → 목차 → 그림 목차 → 표 목차 → 국문초록 → 본문 → 참고문헌 → 영문초록 → 부록",
        "- 제목 번호: 1., 가., 1), 가) 흐름",
        "",
        "## 표 목록",
        *[f"- {item}" for item in builder.table_log],
        "",
        "## 그림 목록",
        *[f"- {item}" for item in builder.figure_log],
        "",
        "## 텍스트 검증",
        f"- 공백 제외 글자 수: {validation['char_count_no_space']}",
        f"- 깨진 문자 수: {validation['replacement_character_count']}",
        f"- API key 패턴 수: {validation['api_key_pattern_count']}",
        f"- 금지 주장 hit: {', '.join(validation['forbidden_phrase_hits']) if validation['forbidden_phrase_hits'] else '없음'}",
        f"- 핵심 수치 포함: {validation['required_metric_hits']}",
        f"- 필수 순서 항목 포함: {validation['required_order_hits']}",
        "",
        "## 최종 편집 메모",
        "- [성명], [학번], [지도교수], [제출일]은 제출 전 직접 입력",
        "- 학교 HWP 원본 양식의 직인/서명란은 한글에서 최종 확인",
        "- 그림/표 번호와 본문 설명 문장 일치 여부 육안 확인",
    ]
    CHECKLIST_PATH.write_text("\n".join(lines), encoding="utf-8-sig")


def main() -> None:
    OUTPUTS.mkdir(exist_ok=True)
    v4.HEADING_FONT = HEADING_FONT
    evidence = v4.load_evidence()
    builder = ThesisV5Builder(evidence)
    builder.add_front_matter()
    builder.add_body()
    builder.save()
    text = extract_docx_text(DOCX_PATH)
    validation = validate_text(text)
    validation.update({
        "docx_path": str(DOCX_PATH),
        "table_count": builder.table_no,
        "figure_count": builder.figure_no,
    })
    VALIDATION_PATH.write_text(json.dumps(validation, ensure_ascii=False, indent=2), encoding="utf-8")
    write_checklist(builder, validation)
    print(f"DOCX created: {DOCX_PATH}")
    print(f"tables={builder.table_no}, figures={builder.figure_no}, chars_no_space={validation['char_count_no_space']}")
    print(f"checklist: {CHECKLIST_PATH}")


if __name__ == "__main__":
    main()
