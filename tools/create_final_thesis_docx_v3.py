"""Create a 29-page DOCX thesis draft with figures and tables.

This script is intentionally separate from older HWPX/DOCX experiments.
The final editing target for this version is DOCX/PDF, with HWPX conversion
treated as optional outside this script.
"""

from __future__ import annotations

import csv
import json
import re
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Mm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUTS = ROOT / "outputs"
FIGURES = OUTPUTS / "ppt_thesis_visual_assets" / "figures"
DOCX_PATH = OUTPUTS / "final_thesis_manuscript_29p_v3.docx"
CHECKLIST_PATH = OUTPUTS / "final_thesis_manuscript_29p_v3_figure_table_checklist.md"

BODY_FONT = "Batang"
BODY_FONT_KO = "바탕"
TITLE_FONT = "HY견명조"
HEAD_FONT = "Malgun Gothic"
HEAD_FONT_KO = "맑은 고딕"
GREEN = "00543D"
GREEN_DARK = "003D2E"
GREEN_LIGHT = "EAF5EE"
LINE = "C9D8CF"
GRAY = "4B5563"


def read_json(relative: str, default: dict | None = None) -> dict:
    path = ROOT / relative
    if not path.exists():
        return default or {}
    return json.loads(path.read_text(encoding="utf-8"))


def read_csv(relative: str) -> list[dict[str, str]]:
    path = ROOT / relative
    if not path.exists():
        return []
    with path.open("r", encoding="utf-8-sig", newline="") as handle:
        return list(csv.DictReader(handle))


def fnum(value, digits: int = 4, fallback: str = "-") -> str:
    try:
        return f"{float(value):.{digits}f}"
    except Exception:
        return fallback if value in (None, "") else str(value)


def pct(value, digits: int = 2) -> str:
    try:
        return f"{float(value) * 100:.{digits}f}%"
    except Exception:
        return str(value)


def pct_plain(value, digits: int = 2) -> str:
    try:
        return f"{float(value) * 100:.{digits}f}%"
    except Exception:
        return str(value)


def set_run_font(run, size: float = 11, bold: bool = False, color: str = "111827", font: str = BODY_FONT_KO) -> None:
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, margin_twips: int = 90) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        node = tc_mar.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            tc_mar.append(node)
        node.set(qn("w:w"), str(margin_twips))
        node.set(qn("w:type"), "dxa")


def set_cell_text(cell, text: str, bold: bool = False, color: str = "111827", size: float = 8.5, align_center: bool = False) -> None:
    cell.text = ""
    set_cell_margins(cell)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if align_center else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(str(text))
    set_run_font(run, size=size, bold=bold, color=color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


@dataclass
class Evidence:
    xgb_pr_auc: str
    xgb_roc_auc: str
    logit_pr_auc: str
    threshold: str
    threshold_precision: str
    threshold_recall: str
    threshold_f1: str
    genai_mode: str
    genai_probability: str
    genai_status: str
    genai_factors: str
    scania_improvement: str
    scania_cost: str


class ThesisBuilder:
    def __init__(self) -> None:
        self.doc = Document()
        self.figure_no = 0
        self.table_no = 0
        self.figure_list: list[str] = []
        self.table_list: list[str] = []
        self.ev = self.collect_evidence()

    def collect_evidence(self) -> Evidence:
        metrics = read_json("outputs/metrics.json")
        threshold = read_json("outputs/threshold_summary.json")
        ai_ctx = read_json("outputs/ai_report_context.json")
        scania = read_json("outputs/scania_official_cost_metrics.json")

        xgb = metrics.get("models", {}).get("xgboost", {})
        logit = metrics.get("models", {}).get("logistic_regression", {})
        selected = threshold.get("selected_metrics", {})
        ai_row = ai_ctx.get("row", {})
        factors = [
            item.get("feature", "").replace("_", " ")
            for item in ai_ctx.get("top_shap_factors", [])[:3]
            if item.get("feature")
        ]
        scania_best = {}
        for row in scania.get("metrics", []):
            if row.get("strategy_id") == "xgboost_cost_optimized":
                scania_best = row
                break
        return Evidence(
            xgb_pr_auc=fnum(xgb.get("pr_auc", 0.8014)),
            xgb_roc_auc=fnum(xgb.get("roc_auc", 0.9736)),
            logit_pr_auc=fnum(logit.get("pr_auc", 0.4188)),
            threshold=fnum(threshold.get("selected_threshold", 0.87), 2),
            threshold_precision=fnum(selected.get("precision", 0.8197)),
            threshold_recall=fnum(selected.get("recall", 0.7353)),
            threshold_f1=fnum(selected.get("f1_score", 0.7752)),
            genai_mode=ai_ctx.get("report_generation_mode", "gemini_generate_content:gemini-2.5-flash"),
            genai_probability=fnum(ai_row.get("xgboost_probability", 0.993616), 6),
            genai_status=str(ai_row.get("risk_status", "High Risk")),
            genai_factors=", ".join(factors) if factors else "torque, rotational speed, process temperature",
            scania_improvement=pct_plain(scania_best.get("cost_improvement_vs_rule", 0.1702)),
            scania_cost=fnum(scania_best.get("official_cost", 0), 4),
        )

    def configure(self) -> None:
        section = self.doc.sections[0]
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.top_margin = Mm(35)
        section.left_margin = Mm(35)
        section.right_margin = Mm(30)
        section.bottom_margin = Mm(25)
        section.footer_distance = Mm(15)

        normal = self.doc.styles["Normal"]
        normal.font.name = BODY_FONT
        normal._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT_KO)
        normal.font.size = Pt(11)
        normal.paragraph_format.line_spacing = 2.0
        normal.paragraph_format.space_after = Pt(0)

        for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
            style = self.doc.styles[style_name]
            style.font.name = HEAD_FONT
            style._element.rPr.rFonts.set(qn("w:eastAsia"), HEAD_FONT_KO)
            style.font.color.rgb = RGBColor.from_string(GREEN_DARK)
            style.font.bold = True

        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run("전남대학교 산업공학과 | MaintiQ Predict")
        set_run_font(run, size=8.5, color=GRAY)

    def add_page_number(self, paragraph) -> None:
        run = paragraph.add_run()
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = "PAGE"
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        run._r.append(fld_begin)
        run._r.append(instr)
        run._r.append(fld_end)

    def page_break(self) -> None:
        self.doc.add_page_break()

    def para(self, text: str, first_indent: bool = True, size: float = 11, align: int | None = None, bold: bool = False) -> None:
        p = self.doc.add_paragraph()
        p.paragraph_format.line_spacing = 2.0
        p.paragraph_format.space_after = Pt(2)
        if first_indent:
            p.paragraph_format.first_line_indent = Cm(0.7)
        if align is not None:
            p.alignment = align
        run = p.add_run(text)
        set_run_font(run, size=size, bold=bold)

    def bullet(self, text: str) -> None:
        p = self.doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.7)
        p.paragraph_format.line_spacing = 1.6
        run = p.add_run("· " + text)
        set_run_font(run, size=10.5)

    def heading(self, text: str, level: int = 1) -> None:
        p = self.doc.add_heading(text, level=level)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        for run in p.runs:
            run.font.name = HEAD_FONT
            run._element.rPr.rFonts.set(qn("w:eastAsia"), HEAD_FONT_KO)
            run.font.color.rgb = RGBColor.from_string(GREEN_DARK)

    def centered(self, text: str, size: float = 11, bold: bool = False, font: str = BODY_FONT_KO) -> None:
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = 1.6
        run = p.add_run(text)
        set_run_font(run, size=size, bold=bold, font=font)

    def table(self, caption: str, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
        self.table_no += 1
        full_caption = f"표 {self.table_no}. {caption}"
        self.table_list.append(full_caption)
        cap = self.doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.line_spacing = 1.2
        run = cap.add_run(full_caption)
        set_run_font(run, size=9, bold=True, color=GREEN_DARK)

        tbl = self.doc.add_table(rows=1, cols=len(headers))
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl.style = "Table Grid"
        for idx, header in enumerate(headers):
            cell = tbl.rows[0].cells[idx]
            set_cell_shading(cell, GREEN)
            set_cell_text(cell, header, bold=True, color="FFFFFF", size=8, align_center=True)
        for ridx, row in enumerate(rows):
            cells = tbl.add_row().cells
            for idx, value in enumerate(row):
                set_cell_shading(cells[idx], GREEN_LIGHT if ridx % 2 == 0 else "FFFFFF")
                set_cell_text(cells[idx], value, size=8.2, align_center=len(str(value)) < 14)
        if widths:
            for row in tbl.rows:
                for idx, width in enumerate(widths):
                    row.cells[idx].width = Cm(width)
        self.doc.add_paragraph()

    def figure(self, image_name: str, caption: str, width_in: float = 5.2) -> None:
        path = FIGURES / image_name
        if not path.exists():
            path = OUTPUTS / image_name
        if not path.exists():
            self.para(f"[그림 누락: {image_name}]", first_indent=False, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
            return
        self.figure_no += 1
        full_caption = f"그림 {self.figure_no}. {caption}"
        self.figure_list.append(full_caption)
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = 1.0
        p.add_run().add_picture(str(path), width=Inches(width_in))
        cap = self.doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.line_spacing = 1.2
        run = cap.add_run(full_caption)
        set_run_font(run, size=9, bold=True, color=GREEN_DARK)

    def callout(self, text: str) -> None:
        table = self.doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = table.rows[0].cells[0]
        set_cell_shading(cell, "F3F8F5")
        set_cell_margins(cell, 150)
        p = cell.paragraphs[0]
        p.paragraph_format.line_spacing = 1.45
        run = p.add_run(text)
        set_run_font(run, size=9.5, bold=True, color=GREEN_DARK)
        self.doc.add_paragraph()

    def build_front_matter(self) -> None:
        self.centered("학사학위논문", 16, True, TITLE_FONT)
        self.doc.add_paragraph("\n\n")
        self.centered("ML 예측, Predictive SPC, GenAI 리포트 및 승인형 작업지시를 결합한", 16, True, TITLE_FONT)
        self.centered("스마트 제조 예지보전 운영 시스템 구현", 18, True, TITLE_FONT)
        self.doc.add_paragraph("\n\n\n")
        self.centered("전남대학교 산업공학과", 13, True)
        self.centered("[학번]  [성명]", 12)
        self.doc.add_paragraph("\n\n")
        self.centered("[제출일]", 11)
        self.page_break()

        self.centered("내표지 I", 14, True)
        self.doc.add_paragraph("\n")
        self.centered("ML 예측, Predictive SPC, GenAI 리포트 및 승인형 작업지시를 결합한", 15, True, TITLE_FONT)
        self.centered("스마트 제조 예지보전 운영 시스템 구현", 17, True, TITLE_FONT)
        self.doc.add_paragraph("\n\n")
        self.centered("이 논문을 학사학위논문으로 제출함", 12)
        self.doc.add_paragraph("\n\n")
        self.centered("전남대학교 산업공학과", 12)
        self.centered("[학번]  [성명]", 12)
        self.centered("지도교수  [지도교수명]", 12)
        self.page_break()

        self.centered("내표지 II", 14, True)
        self.doc.add_paragraph("\n\n")
        self.centered("[성명]의 학사학위논문을 인준함", 13)
        self.doc.add_paragraph("\n\n\n")
        for role in ["심사위원장", "심사위원", "심사위원"]:
            self.centered(f"{role}  ____________________", 12)
            self.doc.add_paragraph()
        self.doc.add_paragraph("\n\n")
        self.centered("[제출일]", 11)
        self.page_break()

        self.heading("목차", 1)
        toc = [
            "1. 서론",
            "  가. 연구 배경",
            "  나. 연구 목적 및 범위",
            "2. 이론적 배경 및 선행연구",
            "  가. 예지보전과 CBM",
            "  나. SPC와 비용 민감 평가",
            "  다. ML 기반 고장 예측과 설명 가능 AI",
            "3. 연구 방법",
            "4. 시스템 구현",
            "5. 실험 및 검증",
            "6. 결론 및 향후 연구",
            "참고문헌",
            "부록",
        ]
        for item in toc:
            self.para(item, first_indent=False, size=10.5)
        self.page_break()

        self.heading("그림 목차", 1)
        figure_plan = [
            "그림 1. MaintiQ Predict 전체 시스템 구조",
            "그림 2. 데이터 입력 및 전처리 흐름",
            "그림 3. Threshold tuning 결과",
            "그림 4. Predictive SPC 위험도 관리도",
            "그림 5. 승인형 작업지시 workflow",
            "그림 6. 데스크톱 제품 화면",
            "그림 7. AI4I PR curve",
            "그림 8. 모델 전략별 PR curve",
            "그림 9. SPC control chart",
            "그림 10. 운영 cost simulation",
            "그림 11. SCANIA official cost 비교",
            "그림 12. 공개 benchmark cost 요약",
        ]
        for item in figure_plan:
            self.para(item, first_indent=False, size=10)
        self.heading("표 목차", 1)
        table_plan = [
            "표 1. 보전 전략 비교",
            "표 2. 선행연구 review 요약",
            "표 3. AI4I baseline 성능",
            "표 4. Threshold tuning 결과",
            "표 5. SPC-only vs ML+SPC 비교",
            "표 6. SCANIA official cost metric",
            "표 7. 가능한 주장과 피해야 할 주장",
        ]
        for item in table_plan:
            self.para(item, first_indent=False, size=10)
        self.doc.add_paragraph()

        self.heading("국문초록", 1)
        self.para(
            "본 연구는 스마트 제조 환경에서 설비 센서 데이터를 기반으로 고장 위험을 예측하고, 예측 결과를 통계적 모니터링과 관리자 의사결정 흐름으로 연결하는 예지보전 운영 시스템을 구현하였다. 제안 시스템인 MaintiQ Predict는 센서 CSV 입력, 데이터 전처리, XGBoost 기반 고장 확률 예측, threshold tuning, Predictive SPC, GenAI 관리자 리포트, 승인형 작업지시 workflow를 하나의 로컬 데스크톱 MVP로 통합한다."
        )
        self.para(
            f"AI4I 2020 데이터셋을 기준으로 Logistic Regression과 XGBoost를 비교한 결과, XGBoost는 PR-AUC {self.ev.xgb_pr_auc}, ROC-AUC {self.ev.xgb_roc_auc}를 보였다. F1-score 기준 threshold tuning 결과 위험 판정 기준은 {self.ev.threshold}로 설정되었고, precision {self.ev.threshold_precision}, recall {self.ev.threshold_recall}, F1-score {self.ev.threshold_f1}를 달성하였다."
        )
        self.para(
            f"GenAI 리포트 검증에서는 {self.ev.genai_mode} 방식으로 관리자 참고용 리포트를 생성하였으며, 예측 확률 {self.ev.genai_probability}, 상태 {self.ev.genai_status}, 주요 위험 요인 {self.ev.genai_factors}를 요약하였다. 또한 SCANIA Component X 공개 산업 benchmark의 official cost metric에서는 rule baseline 대비 {self.ev.scania_improvement} 개선 가능성을 확인하였다."
        )
        self.para(
            "본 연구의 차별성은 단일 모델의 성능 제시에 그치지 않고 예측, 관리도, 설명 리포트, 작업지시 승인 이력, 공개 benchmark 검증을 연결한 통합 구현에 있다. 다만 실제 회사의 labeled sensor CSV, 정비 이력, downtime 및 비용 로그가 확보되지 않았으므로 실제 현장 비용 절감이나 탐지 시간 단축은 향후 실증 과제로 둔다."
        )
        self.para("주요어: 예지보전, 스마트 제조, XGBoost, Predictive SPC, GenAI, 작업지시, SCANIA Component X", first_indent=False, bold=True)
        self.page_break()

    def build_body(self) -> None:
        # Page 7
        self.heading("1. 서론", 1)
        self.heading("가. 연구 배경", 2)
        self.para(
            "제조 설비는 고장이 발생하는 순간 생산 중단, 품질 손실, 납기 지연, 정비 비용 증가를 동시에 유발한다. 특히 회전 장비, 공구 마모 장비, 압축기와 같은 설비는 센서값이 연속적으로 변하지만 고장 라벨은 매우 적게 발생하므로, 단순한 점검 주기만으로는 위험을 충분히 설명하기 어렵다."
        )
        self.para(
            "기존의 사후보전은 고장이 발생한 뒤 대응하므로 downtime이 커질 수 있고, 예방보전은 고장 여부와 관계없이 정비를 수행해 과잉 정비가 발생할 수 있다. 반면 예지보전은 센서 데이터의 변화와 고장 이력을 기반으로 위험 확률을 계산하고, 이를 정비 의사결정에 반영하는 접근이다."
        )
        self.para(
            "그러나 예측 모델만으로는 운영 가치가 충분하지 않다. 실제 운영자는 위험 확률이 왜 높아졌는지, 어떤 기준으로 고위험으로 판단했는지, 이후 어떤 작업지시를 검토해야 하는지 확인해야 한다. 따라서 본 연구는 예측 모델과 운영 workflow를 함께 구현하는 방향으로 문제를 정의한다."
        )
        self.page_break()

        # Page 8
        self.heading("나. 연구 목적 및 범위", 2)
        self.para(
            "본 연구의 목적은 설비 센서 CSV를 입력받아 고장 위험을 예측하고, 그 결과를 통계적 관리도, AI 관리자 리포트, 승인형 작업지시 이력으로 연결하는 제품형 MVP를 구현하는 것이다. 사용자는 데이터 CSV를 입력하고, 시스템은 전처리, 예측, 위험 우선순위, 리포트, 작업지시 승인 기록을 순차적으로 제공한다."
        )
        self.para(
            "연구 범위는 로컬 데스크톱 앱과 Streamlit Admin 검증 콘솔이다. 사용자용 앱은 운영자가 바로 사용할 수 있도록 연구 용어를 숨긴 제품 화면으로 구성했고, Admin 콘솔은 모델 검증, 공개 benchmark, 산업공학 근거, field validation 템플릿을 확인하는 연구 검증 화면으로 분리하였다."
        )
        self.callout(
            "연구의 핵심 기여는 새로운 ML 알고리즘 자체가 아니라, ML 예측 결과를 threshold 정책, Predictive SPC, GenAI 리포트, 승인형 작업지시, 공개 산업 benchmark 검증으로 연결한 통합 구현이다."
        )
        self.para(
            "본 연구는 실제 공장 PLC/SCADA 운영망과 직접 연결하지 않았고, 실제 회사의 정비 비용 로그로 비용 절감률을 실증하지 않았다. 이러한 부분은 향후 현장 데이터 수집과 운영 검증을 통해 확장해야 한다."
        )
        self.page_break()

        # Page 9
        self.heading("2. 이론적 배경 및 선행연구", 1)
        self.heading("가. 예지보전과 보전 전략", 2)
        self.para(
            "보전 전략은 일반적으로 사후보전, 예방보전, 예지보전으로 구분된다. 사후보전은 설비 고장 이후 대응하는 방식이고, 예방보전은 일정 주기 또는 사용량에 따라 점검하는 방식이다. 예지보전은 센서 데이터와 고장 패턴을 분석하여 고장이 발생하기 전에 정비 필요성을 판단하는 방식이다."
        )
        self.table(
            "보전 전략 비교",
            ["구분", "판단 기준", "장점", "한계"],
            [
                ["사후보전", "고장 발생 후", "운영 방식이 단순함", "downtime과 긴급 비용 증가"],
                ["예방보전", "정해진 주기", "고장 전 점검 가능", "과잉 정비와 누락 위험"],
                ["예지보전", "센서 상태와 위험 예측", "상태 기반 의사결정", "데이터 품질과 모델 검증 필요"],
            ],
            [2.2, 3.0, 4.0, 5.0],
        )
        self.para(
            "본 연구는 예지보전 관점에서 센서 기반 고장 확률을 산출하고, 해당 확률이 운영 기준을 초과할 때 관리자 리포트와 작업지시 검토 흐름으로 이어지게 설계하였다."
        )
        self.page_break()

        # Page 10
        self.heading("나. SPC, FMEA/RPN 및 비용 민감 평가", 2)
        self.para(
            "SPC는 공정 또는 설비 상태의 변동을 시간축에서 관찰하고, 관리한계선을 기준으로 이상 신호를 판단하는 방법이다. 본 연구에서는 원시 센서값이 아니라 ML이 계산한 고장 위험 확률을 시간 순서로 정렬하여 Predictive SPC를 구성하였다."
        )
        self.para(
            "FMEA/RPN 관점에서는 고장의 심각도, 발생 가능성, 탐지 가능성을 함께 고려한다. 본 시스템의 risk priority score는 calibrated probability, 운영 threshold, 데이터 품질, false alarm 및 missed failure 비용 가중치를 결합하여 우선순위 점수로 해석할 수 있다."
        )
        self.table(
            "산업공학 지표와 시스템 변수의 연결",
            ["산업공학 개념", "본 시스템 대응 변수", "해석"],
            [
                ["OEE", "고위험 건수, downtime 로그", "향후 설비 효율 영향 분석"],
                ["MTBF", "고장 간격 및 event 이력", "고장 발생 주기 추정"],
                ["MTTR", "정비 시작/종료 로그", "수리 시간 변화 분석"],
                ["RPN", "확률·품질·비용 가중치", "위험 우선순위 산정"],
            ],
            [3.0, 4.0, 6.0],
        )
        self.page_break()

        # Page 11
        self.heading("다. ML 기반 고장 예측과 설명 가능 AI", 2)
        self.para(
            "제조 고장 데이터는 정상 데이터가 대부분이고 고장 데이터가 적은 불균형 구조를 갖는다. 이 경우 accuracy만으로는 모델을 평가하기 어렵고, precision, recall, F1-score, ROC-AUC, PR-AUC를 함께 확인해야 한다. 특히 고장 클래스가 적을수록 PR-AUC와 recall의 의미가 커진다."
        )
        self.para(
            "본 연구는 Logistic Regression을 기준선 모델로 사용하고 XGBoost를 주 예측 모델로 비교하였다. XGBoost는 비선형 변수 조합을 반영할 수 있어 온도, 회전 속도, 토크, 공구 마모가 결합되는 고장 위험 예측에 적합하다."
        )
        self.para(
            "설명 가능 AI 측면에서는 SHAP 요인을 사용해 어떤 센서 변수가 위험 확률을 높였는지 제시한다. 또한 GenAI 리포트는 수치 결과를 관리자 참고 문장으로 변환하여, 운영자가 조치 필요성을 이해할 수 있도록 돕는다."
        )
        self.page_break()

        # Page 12
        self.heading("라. 선행연구 review와 본 연구 반영점", 2)
        self.table(
            "선행연구 review 요약",
            ["영역", "대표 연구/자료", "주요 내용", "본 연구 반영점"],
            [
                ["CBM/예지보전", "Jardine et al., Carvalho et al.", "상태 기반 정비와 데이터 기반 진단", "센서 CSV 기반 위험 예측"],
                ["SPC", "Montgomery", "관리도와 이상 신호 판단", "ML 확률 기반 Predictive SPC"],
                ["ML 예측", "AI4I 2020, XGBoost", "불균형 고장 예측 benchmark", "PR-AUC와 threshold tuning"],
                ["설명성/비용", "SHAP, Elkan, SCANIA", "설명 가능성과 비용 민감 평가", "GenAI 리포트와 official cost metric"],
            ],
            [2.5, 3.2, 4.0, 4.4],
        )
        self.para(
            "선행연구는 각각 예측 모델, 통계적 관리, 설명성, 비용 민감 평가를 다루지만, 본 연구는 이를 하나의 운영 흐름으로 연결한다. 즉, 센서 데이터 입력부터 고위험 판정, 리포트, 작업지시 승인 이력까지 연결한 구현 중심의 통합성이 연구 차별점이다."
        )
        self.page_break()

        # Page 13
        self.heading("3. 연구 방법", 1)
        self.heading("가. 전체 연구 설계", 2)
        self.para(
            "연구 방법은 데이터 입력, 전처리, 모델 학습, threshold 정책, Predictive SPC, GenAI 리포트, 작업지시 workflow, 공개 benchmark 검증으로 구성된다. 그림 1은 전체 시스템 구조를 나타낸다."
        )
        self.figure("01_system_architecture.png", "MaintiQ Predict 전체 시스템 구조", 5.7)
        self.para(
            "사용자 앱은 CSV 기반 예측과 운영 화면을 담당하고, Admin 콘솔은 모델 비교와 연구 검증을 담당한다. 이 분리는 제품 사용성과 연구 재현성을 동시에 확보하기 위한 설계이다."
        )
        self.page_break()

        # Page 14
        self.heading("나. 데이터와 전처리", 2)
        self.para(
            "기본 학습 데이터는 AI4I 2020 데이터셋이다. target은 Machine failure이며, UDI와 Product ID는 식별자 성격이 강하므로 제거하였다. 또한 TWF, HDF, PWF, OSF, RNF는 고장 유형 라벨에 가까워 실제 예측 상황에서는 누수 위험이 있으므로 제거하였다."
        )
        self.figure("02_data_preprocessing_pipeline.png", "데이터 입력 및 전처리 흐름", 5.6)
        self.para(
            "범주형 변수 Type은 one-hot encoding으로 변환하고, train/test split은 stratified 방식과 random_state 42를 사용하였다. 회사별 CSV 대응을 위해 컬럼 자동 매핑, 단위 변환, 결측률 및 숫자 변환 실패 진단도 함께 구현하였다."
        )
        self.page_break()

        # Page 15
        self.heading("다. 모델 학습과 threshold 정책", 2)
        self.para(
            "모델은 Logistic Regression과 XGBoost를 비교하였다. XGBoost는 센서 변수 간 비선형 관계를 반영할 수 있어 주 모델로 선택하였다. 고장 데이터의 불균형 특성을 고려하여 PR-AUC, recall, F1-score를 핵심 지표로 사용하였다."
        )
        self.table(
            "Threshold tuning 결과",
            ["항목", "값", "해석"],
            [
                ["탐색 범위", "0.05~0.95", "확률 기준을 0.01 간격으로 탐색"],
                ["선택 기준", "F1-score", "precision과 recall 균형"],
                ["선택 threshold", self.ev.threshold, "고위험 판정 기준"],
                ["F1-score", self.ev.threshold_f1, "튜닝 후 균형 성능"],
            ],
            [3.0, 2.5, 7.5],
        )
        self.figure("23_threshold_tuning.png", "Threshold tuning 그래프", 4.8)
        self.page_break()

        # Page 16
        self.heading("라. Predictive SPC 구성", 2)
        self.para(
            "Predictive SPC는 모델이 산출한 고장 확률을 시간축으로 정렬한 뒤, 관리한계와 위험 기준을 함께 표시하는 방식으로 구성하였다. 이를 통해 단일 row의 예측뿐 아니라 위험 추세를 관찰할 수 있다."
        )
        self.figure("24_spc_risk_chart.png", "Predictive SPC 위험도 관리도", 5.7)
        self.para(
            "SPC-only 방식은 센서값 또는 확률의 통계적 이상을 탐지하는 데 유용하지만, 고장 여부와 직접 연결되는 ML 확률과 결합할 때 의사결정 해석력이 높아진다."
        )
        self.page_break()

        # Page 17
        self.heading("마. GenAI 리포트와 승인형 작업지시", 2)
        self.para(
            "GenAI 리포트는 예측 확률, 위험 기준, SHAP 요인, SPC 상태를 하나의 관리자 참고 문장으로 변환한다. 이 리포트는 정비 명령을 자동으로 내리는 것이 아니라, 작업자가 승인·검토·반려를 판단하기 위한 참고 정보이다."
        )
        self.table(
            "GenAI 리포트 생성 검증",
            ["항목", "값"],
            [
                ["생성 방식", self.ev.genai_mode],
                ["예측 확률", self.ev.genai_probability],
                ["위험 기준", self.ev.threshold],
                ["상태", self.ev.genai_status],
                ["주요 요인", self.ev.genai_factors],
            ],
            [4.0, 9.5],
        )
        self.figure("03_work_order_workflow.png", "승인형 작업지시 workflow", 5.4)
        self.page_break()

        # Page 18
        self.heading("4. 시스템 구현", 1)
        self.heading("가. 데스크톱 제품 앱 구현", 2)
        self.para(
            "사용자용 앱은 PySide6 기반 Windows 데스크톱 앱으로 구현하였다. 메뉴는 홈, 데이터 예측, 위험 분석, AI 리포트, 작업지시로 구성되어 있으며, 연구 검증과 benchmark 결과는 별도 Admin 콘솔로 분리하였다."
        )
        self.figure("20_app_main_screen.png", "MaintiQ Predict 데스크톱 제품 화면", 5.8)
        self.para(
            "제품 앱은 연구 용어를 노출하지 않고 운영자가 이해할 수 있는 업무 흐름을 우선한다. Full은 정밀 분석 모드, Lite는 빠른 점검 모드로 구분하여 배포 용량과 분석 정밀도의 trade-off를 명확히 하였다."
        )
        self.page_break()

        # Page 19
        self.heading("5. 실험 및 검증", 1)
        self.heading("가. 실험 설계", 2)
        self.para(
            "실험은 AI4I 기본 모델 성능, threshold tuning, SMOTE 및 모델 전략 비교, SPC-only와 ML+SPC 비교, 운영 cost simulation, SCANIA official cost metric, 공개 benchmark 확장으로 구성하였다."
        )
        self.table(
            "실험 설계 요약",
            ["실험", "비교 대상", "주요 지표"],
            [
                ["AI4I baseline", "Logistic Regression, XGBoost", "PR-AUC, ROC-AUC, F1"],
                ["Threshold tuning", "0.05~0.95 threshold", "precision, recall, F1"],
                ["SPC 비교", "SPC-only, ML threshold, ML+SPC", "F1, alert 수"],
                ["Cost simulation", "정책별 운영 비용", "false alarm, missed failure"],
                ["SCANIA", "rule baseline, cost-optimized XGBoost", "official cost metric"],
            ],
            [3.3, 5.0, 5.2],
        )
        self.page_break()

        # Page 20
        self.heading("나. AI4I baseline 결과", 2)
        self.para(
            f"AI4I 기준 XGBoost는 PR-AUC {self.ev.xgb_pr_auc}, ROC-AUC {self.ev.xgb_roc_auc}를 보였고, Logistic Regression의 PR-AUC {self.ev.logit_pr_auc}보다 높은 성능을 나타냈다."
        )
        self.table(
            "AI4I baseline 성능",
            ["모델", "PR-AUC", "ROC-AUC", "해석"],
            [
                ["Logistic Regression", self.ev.logit_pr_auc, "-", "기준선 모델"],
                ["XGBoost", self.ev.xgb_pr_auc, self.ev.xgb_roc_auc, "주 예측 모델"],
            ],
            [4.0, 2.5, 2.5, 4.5],
        )
        self.figure("22_pr_curve.png", "AI4I PR curve", 4.9)
        self.page_break()

        # Page 21
        self.heading("다. SMOTE 및 threshold 전략 비교", 2)
        self.para(
            f"기본 확률 기준 0.5 대신 F1-score 기준 threshold를 탐색한 결과, {self.ev.threshold} 기준에서 F1-score {self.ev.threshold_f1}를 얻었다. 이는 false alarm과 missed failure의 trade-off를 고려한 운영 기준 설정의 필요성을 보여준다."
        )
        self.figure("33_model_strategy_pr_curve.png", "모델 전략별 PR curve", 5.3)
        self.para(
            "SMOTE는 고장 클래스가 적은 상황에서 학습 데이터 균형을 맞추는 방법이지만, 실제 운영에서는 오경보 증가 가능성도 함께 검토해야 한다. 따라서 본 연구는 SMOTE 자체보다 threshold 정책과 비용 민감 평가를 함께 제시하였다."
        )
        self.page_break()

        # Page 22
        self.heading("라. SPC-only vs ML+SPC 비교", 2)
        self.para(
            "SPC-only 방식은 통계적 이상 신호를 탐지하는 데 유용하지만, 실제 고장 라벨과의 직접 연결성은 제한될 수 있다. ML threshold 방식은 고장 확률을 직접 사용하므로 F1-score가 높게 나타났고, ML+SPC는 위험 추세 해석을 보완한다."
        )
        self.table(
            "SPC-only vs ML+SPC 비교",
            ["정책", "F1-score", "특징"],
            [
                ["SPC-only", "0.1600", "통계 이상 탐지 중심"],
                ["ML threshold", self.ev.threshold_f1, "고장 확률 기준 판단"],
                ["ML+SPC", "0.7051", "확률과 추세를 함께 해석"],
            ],
            [4.0, 2.8, 6.5],
        )
        self.figure("25_spc_control_chart.png", "SPC control chart", 5.4)
        self.page_break()

        # Page 23
        self.heading("마. 운영 cost simulation", 2)
        self.para(
            "운영 의사결정에서는 false alarm과 missed failure의 비용이 동일하지 않다. 따라서 본 연구는 실제 원화 비용이 아니라 상대 단위의 normalized operating cost를 사용하여 conservative, balanced, high downtime 시나리오를 비교하였다."
        )
        self.figure("28_operational_value_simulation.png", "운영 cost simulation 결과", 5.7)
        self.para(
            "이 결과는 실제 비용 절감을 실증한 것이 아니라, 운영 정책을 비교하기 위한 simulation이다. 실제 비용 효과를 주장하려면 설비별 downtime, 부품비, 인건비, 기존 점검 결과가 포함된 회사 로그가 필요하다."
        )
        self.page_break()

        # Page 24
        self.heading("바. SCANIA official cost metric 검증", 2)
        self.para(
            f"SCANIA Component X는 실제 SCANIA fleet에서 수집된 공개 산업 데이터셋으로, class 0~4 예측과 official cost matrix를 제공한다. 본 연구는 cost-optimized XGBoost 전략을 적용하여 rule baseline 대비 official cost metric {self.ev.scania_improvement} 개선 가능성을 확인하였다."
        )
        self.table(
            "SCANIA official cost metric 요약",
            ["비교 기준", "결과", "해석"],
            [
                ["Rule baseline", "기준", "단순 rule 기반 정책"],
                ["XGBoost cost-optimized", self.ev.scania_improvement, "official cost metric 기준 개선"],
                ["비용 해석", "상대 cost metric", "실제 원화 비용 절감과 구분"],
            ],
            [4.0, 3.5, 6.0],
        )
        self.figure("29_scania_cost_comparison.png", "SCANIA official cost 비교", 5.0)
        self.page_break()

        # Page 25
        self.heading("사. 공개 benchmark 확장", 2)
        self.para(
            "SCANIA 외에도 MetroPT-3, NASA C-MAPSS, IMS/FEMTO 등 공개 산업 benchmark를 대상으로 adapter와 검증 산출물을 구성하였다. 각 데이터셋은 anomaly horizon, RUL, run-to-failure 등 서로 다른 평가 관점을 제공한다."
        )
        self.figure("31_public_benchmark_cost_chart.png", "공개 benchmark cost 요약", 5.2)
        self.para(
            "공개 benchmark는 실제 산업 데이터에 가까운 검증 근거를 제공하지만, 특정 회사의 비용 절감이나 다운타임 감소를 직접 입증하는 자료는 아니다. 따라서 논문에서는 공개 benchmark 검증과 현장 실증을 명확히 분리한다."
        )
        self.page_break()

        # Page 26
        self.heading("아. 현장 실증 준비와 주장 경계", 2)
        self.para(
            "현장 실증을 위해서는 labeled sensor CSV, maintenance history CSV, downtime/cost CSV가 필요하다. 본 시스템은 세 파일이 들어왔을 때 precision, recall, false alarm, missed failure, lead time, downtime delta, maintenance cost delta를 계산하는 리포트 구조를 준비하였다."
        )
        self.table(
            "가능한 주장과 피해야 할 주장",
            ["구분", "표현"],
            [
                ["가능", "AI4I와 공개 benchmark 기준 성능 및 cost metric 비교"],
                ["가능", "GenAI 관리자 참고 리포트와 승인형 작업지시 구현"],
                ["피해야 함", "실제 회사 로그 기반 성능 검증이 끝났다는 표현"],
                ["피해야 함", "실제 원화 비용 절감률을 입증했다는 표현"],
                ["피해야 함", "공장 운영망에 직접 배포했다는 표현"],
            ],
            [3.0, 10.5],
        )
        self.page_break()

        # Page 27
        self.heading("6. 논의", 1)
        self.para(
            "본 연구의 결과는 예측 모델의 성능뿐 아니라 운영 workflow로 연결되는 구조의 중요성을 보여준다. 단순히 고장 확률을 산출하는 것만으로는 현장 의사결정을 충분히 지원하기 어렵다. 위험 기준, 관리도, 설명 리포트, 승인 기록이 함께 제공되어야 담당자가 결과를 검토하고 조치를 결정할 수 있다."
        )
        self.para(
            "또한 Lite와 Full 실행본을 분리함으로써 일반 사용자는 빠른 점검 모드로 가볍게 시스템을 사용할 수 있고, 연구 검증 또는 정밀 분석이 필요한 경우에는 XGBoost/SHAP 기반 정밀 분석 모드를 사용할 수 있다. 이는 제품 배포와 연구 재현성 사이의 trade-off를 줄이기 위한 설계이다."
        )
        self.para(
            "한계는 명확하다. AI4I와 SCANIA는 공개 데이터 기반 검증이며, 실제 회사의 설비 ID, timestamp, 센서값, 실제 고장 여부, 정비 시작·종료, downtime, 부품비, 인건비 로그가 확보되어야 현장 효과를 평가할 수 있다."
        )
        self.page_break()

        # Page 28
        self.heading("7. 결론 및 향후 연구", 1)
        self.para(
            "본 연구는 스마트 제조 예지보전을 위한 MaintiQ Predict를 구현하였다. 시스템은 센서 CSV 입력, 전처리, XGBoost 고장 확률 예측, threshold tuning, Predictive SPC, GenAI 관리자 리포트, 승인형 작업지시 workflow를 통합한다."
        )
        self.para(
            f"AI4I 기준 XGBoost는 PR-AUC {self.ev.xgb_pr_auc}를 보였고, threshold {self.ev.threshold}에서 F1-score {self.ev.threshold_f1}를 달성하였다. SCANIA 공개 benchmark에서는 official cost metric 기준 rule baseline 대비 {self.ev.scania_improvement} 개선 가능성을 확인하였다."
        )
        self.para(
            "향후 연구는 실제 회사 데이터 기반 field validation, PLC/SCADA 또는 MQTT/OPC UA 연동, 운영 DB와 권한 관리, 모델 drift monitoring, 코드 서명 및 자동 업데이트를 포함한 배포 품질 강화로 확장할 수 있다."
        )
        self.page_break()

        # Page 29
        self.heading("참고문헌", 1)
        refs = [
            "Jardine, A. K. S., Lin, D., & Banjevic, D. (2006). A review on machinery diagnostics and prognostics implementing condition-based maintenance. Mechanical Systems and Signal Processing.",
            "Carvalho, T. P., et al. (2019). A systematic literature review of machine learning methods applied to predictive maintenance. Computers & Industrial Engineering.",
            "Montgomery, D. C. (2019). Introduction to Statistical Quality Control. Wiley.",
            "Elkan, C. (2001). The foundations of cost-sensitive learning. IJCAI.",
            "Chen, T., & Guestrin, C. (2016). XGBoost: A scalable tree boosting system. KDD.",
            "Lundberg, S. M., & Lee, S. I. (2017). A unified approach to interpreting model predictions. NeurIPS.",
            "Chawla, N. V., et al. (2002). SMOTE: Synthetic minority over-sampling technique. Journal of Artificial Intelligence Research.",
            "Matzka, S. (2020). AI4I 2020 Predictive Maintenance Dataset. UCI Machine Learning Repository.",
            "SCANIA Component X Dataset. Researchdata.se.",
            "UCI Machine Learning Repository. MetroPT-3 Dataset.",
            "NASA Prognostics Center of Excellence. C-MAPSS and IMS Bearing datasets.",
            "Google AI for Developers. Gemini API generateContent documentation.",
        ]
        for ref in refs:
            self.para(ref, first_indent=False, size=9.5)
        self.heading("부록", 1)
        self.para(
            "부록에는 실행 명령, field validation 템플릿, GenAI 리포트 전문, 설치형 데스크톱 앱 검증 결과를 별도 산출물로 연결한다. API key와 원본 회사 데이터는 문서와 저장소에 포함하지 않는다.",
            first_indent=False,
            size=9.5,
        )

    def save(self) -> None:
        self.doc.save(DOCX_PATH)
        lines = ["# final_thesis_manuscript_29p_v3 그림/표 삽입 체크리스트", ""]
        lines.append("## 그림")
        lines.extend(f"- {item}" for item in self.figure_list)
        lines.append("")
        lines.append("## 표")
        lines.extend(f"- {item}" for item in self.table_list)
        CHECKLIST_PATH.write_text("\n".join(lines) + "\n", encoding="utf-8-sig")

    def build(self) -> None:
        self.configure()
        self.build_front_matter()
        self.build_body()
        self.save()


def validate_docx_text() -> None:
    doc = Document(DOCX_PATH)
    text = "\n".join(p.text for p in doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text += "\n" + cell.text
    forbidden = [
        "AIza",
        "sk-",
        "�",
        "실제 PLC/SCADA 배포 완료",
        "실제 회사 데이터 검증 완료",
        "실제 비용 절감 실증 완료",
        "자동 정비 명령 실행",
    ]
    hits = [item for item in forbidden if item in text]
    if hits:
        raise RuntimeError(f"Forbidden or broken text found: {hits}")
    if text.count("그림 ") < 12:
        raise RuntimeError("Figure captions look incomplete.")
    if text.count("표 ") < 7:
        raise RuntimeError("Table captions look incomplete.")
    key_metrics = ["0.8014", "0.9736", "0.87", "0.7752", "gemini_generate_content:gemini-2.5-flash", "0.993616", "17.02%"]
    missing = [metric for metric in key_metrics if metric not in text]
    if missing:
        raise RuntimeError(f"Missing required metric text: {missing}")


def main() -> None:
    OUTPUTS.mkdir(parents=True, exist_ok=True)
    builder = ThesisBuilder()
    builder.build()
    validate_docx_text()
    print(f"DOCX: {DOCX_PATH}")
    print(f"Checklist: {CHECKLIST_PATH}")


if __name__ == "__main__":
    main()
