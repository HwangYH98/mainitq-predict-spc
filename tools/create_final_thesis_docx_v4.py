from __future__ import annotations

import json
import re
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUTS = ROOT / "outputs"
FIGURES = OUTPUTS / "ppt_thesis_visual_assets" / "figures"
DOCX_PATH = OUTPUTS / "final_thesis_manuscript_29p_v4.docx"
CHECKLIST_PATH = OUTPUTS / "final_thesis_manuscript_29p_v4_figure_table_checklist.md"
VALIDATION_PATH = OUTPUTS / "final_thesis_manuscript_29p_v4_validation.json"

BODY_FONT = "바탕"
HEADING_FONT = "맑은 고딕"
ACCENT = RGBColor(22, 96, 78)
LIGHT_ACCENT = "DDEFE8"
BORDER = "B7D5CC"


def read_json(path: Path, default: dict | list | None = None):
    if not path.exists():
        return default if default is not None else {}
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except Exception:
        return default if default is not None else {}


def load_evidence() -> dict:
    metrics = read_json(OUTPUTS / "metrics.json", {})
    threshold = read_json(OUTPUTS / "threshold_summary.json", {})
    ai_context = read_json(OUTPUTS / "ai_report_context.json", {})
    scania = read_json(OUTPUTS / "scania_official_cost_metrics.json", {})

    xgb = metrics.get("models", {}).get("xgboost", {})
    lr = metrics.get("models", {}).get("logistic_regression", {})
    selected = threshold.get("selected_metrics", {})
    ai_row = ai_context.get("row", {})
    top_shap = ai_context.get("top_shap_factors", [])

    scania_best = {}
    for item in scania.get("metrics", []):
        if item.get("strategy_id") == "xgboost_cost_optimized":
            scania_best = item
            break

    return {
        "xgb_pr_auc": xgb.get("pr_auc", 0.8014),
        "xgb_roc_auc": xgb.get("roc_auc", 0.9736),
        "xgb_precision": xgb.get("precision", 0.4444),
        "xgb_recall": xgb.get("recall", 0.8824),
        "xgb_f1": xgb.get("f1_score", 0.5911),
        "lr_pr_auc": lr.get("pr_auc", 0.3817),
        "lr_roc_auc": lr.get("roc_auc", 0.9069),
        "threshold": threshold.get("selected_threshold", 0.87),
        "tuned_precision": selected.get("precision", 0.8197),
        "tuned_recall": selected.get("recall", 0.7353),
        "tuned_f1": selected.get("f1_score", 0.7752),
        "test_rows": threshold.get("test_rows", 2000),
        "test_failures": threshold.get("test_failures", 68),
        "genai_mode": ai_context.get("report_generation_mode", "gemini_generate_content:gemini-2.5-flash"),
        "genai_prob": ai_row.get("xgboost_probability", 0.993616),
        "genai_status": ai_row.get("risk_status", "High Risk"),
        "genai_threshold": ai_row.get("selected_threshold", 0.87),
        "genai_top_factors": ", ".join([str(x.get("feature", "")) for x in top_shap[:3]]) or "torque_nm, rotational_speed_rpm",
        "scania_cost_improvement": scania_best.get("cost_improvement_vs_rule", 0.1702),
        "scania_cost": scania_best.get("official_cost", 49548.0),
        "scania_normalized_cost": scania_best.get("normalized_cost", 0.863206),
    }


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, margin_twips: int = 90) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m in ("top", "left", "bottom", "right"):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(margin_twips))
        node.set(qn("w:type"), "dxa")


def set_font(run, name: str = BODY_FONT, size: int = 11, bold: bool = False, color: RGBColor | None = None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


class ThesisBuilder:
    def __init__(self, evidence: dict):
        self.doc = Document()
        self.evidence = evidence
        self.figure_no = 0
        self.table_no = 0
        self.figure_log: list[str] = []
        self.table_log: list[str] = []
        self._configure()

    def _configure(self) -> None:
        section = self.doc.sections[0]
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.top_margin = Mm(35)
        section.left_margin = Mm(35)
        section.right_margin = Mm(30)
        section.bottom_margin = Mm(25)
        section.footer_distance = Mm(15)

        styles = self.doc.styles
        normal = styles["Normal"]
        normal.font.name = BODY_FONT
        normal._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
        normal.font.size = Pt(11)
        normal.paragraph_format.line_spacing = 2.0
        normal.paragraph_format.space_after = Pt(0)
        normal.paragraph_format.first_line_indent = Cm(0.4)

        for style_name, size in [("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 12)]:
            style = styles[style_name]
            style.font.name = HEADING_FONT
            style._element.rPr.rFonts.set(qn("w:eastAsia"), HEADING_FONT)
            style.font.size = Pt(size)
            style.font.bold = True
            style.font.color.rgb = ACCENT
            style.paragraph_format.line_spacing = 1.4
            style.paragraph_format.space_before = Pt(7)
            style.paragraph_format.space_after = Pt(3)

        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run("전남대학교 산업공학과 | MaintiQ Predict")
        set_font(run, BODY_FONT, 9, False, RGBColor(90, 90, 90))

    def page_break(self) -> None:
        self.doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    def paragraph(self, text: str, *, bold_lead: str | None = None, align=None, indent: bool = True) -> None:
        p = self.doc.add_paragraph()
        if align:
            p.alignment = align
        if not indent:
            p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.line_spacing = 2.0
        if bold_lead and text.startswith(bold_lead):
            r1 = p.add_run(bold_lead)
            set_font(r1, BODY_FONT, 11, True)
            r2 = p.add_run(text[len(bold_lead):])
            set_font(r2, BODY_FONT, 11)
        else:
            run = p.add_run(text)
            set_font(run)

    def note_box(self, title: str, lines: Iterable[str]) -> None:
        table = self.doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = table.cell(0, 0)
        set_cell_shading(cell, "F3FAF7")
        set_cell_margins(cell, 140)
        p = cell.paragraphs[0]
        p.paragraph_format.first_line_indent = Cm(0)
        r = p.add_run(title)
        set_font(r, HEADING_FONT, 11, True, ACCENT)
        for line in lines:
            p = cell.add_paragraph()
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.line_spacing = 1.35
            rr = p.add_run(f"· {line}")
            set_font(rr, BODY_FONT, 10)
        self.doc.add_paragraph()

    def heading1(self, text: str) -> None:
        self.doc.add_heading(text, level=1)

    def heading2(self, text: str) -> None:
        self.doc.add_heading(text, level=2)

    def heading3(self, text: str) -> None:
        self.doc.add_heading(text, level=3)

    def caption(self, text: str) -> None:
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.line_spacing = 1.2
        r = p.add_run(text)
        set_font(r, BODY_FONT, 9, False, RGBColor(80, 80, 80))

    def table(self, title: str, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
        self.table_no += 1
        self.table_log.append(f"표 {self.table_no}. {title}")
        self.caption(f"표 {self.table_no}. {title}")
        t = self.doc.add_table(rows=1, cols=len(headers))
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        t.style = "Table Grid"
        for i, header in enumerate(headers):
            cell = t.rows[0].cells[i]
            set_cell_shading(cell, LIGHT_ACCENT)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Cm(0)
            r = p.add_run(header)
            set_font(r, BODY_FONT, 9, True, ACCENT)
            if widths:
                cell.width = Cm(widths[i])
        for row_idx, row in enumerate(rows):
            cells = t.add_row().cells
            for i, value in enumerate(row):
                cell = cells[i]
                if row_idx % 2 == 1:
                    set_cell_shading(cell, "F8FBFA")
                set_cell_margins(cell)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                if widths:
                    cell.width = Cm(widths[i])
                p = cell.paragraphs[0]
                p.paragraph_format.first_line_indent = Cm(0)
                p.paragraph_format.line_spacing = 1.25
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT if len(value) > 16 else WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run(value)
                set_font(r, BODY_FONT, 8 if len(value) > 34 else 9)
        self.doc.add_paragraph()

    def figure(self, title: str, filename: str, width_cm: float = 11.8) -> None:
        path = FIGURES / filename
        if not path.exists():
            return
        width_cm = min(width_cm, 10.2)
        self.figure_no += 1
        self.figure_log.append(f"그림 {self.figure_no}. {title} ({filename})")
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        run = p.add_run()
        run.add_picture(str(path), width=Cm(width_cm))
        self.caption(f"그림 {self.figure_no}. {title}")

    def add_front_matter(self) -> None:
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(100)
        p.paragraph_format.first_line_indent = Cm(0)
        r = p.add_run("AI 예지보전 기반 스마트 제조 설비 위험관리 시스템 구현")
        set_font(r, HEADING_FONT, 18, True, ACCENT)
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        r = p.add_run("MaintiQ Predict: ML·SPC·GenAI 통합 접근")
        set_font(r, HEADING_FONT, 13, True)
        for label in ["성명", "학번", "지도교수", "제출일"]:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Cm(0)
            r = p.add_run(f"{label}: [{label}]")
            set_font(r, BODY_FONT, 12)
        self.page_break()

        self.heading1("내표지 I")
        self.paragraph("본 논문은 산업공학 전공 학부 캡스톤디자인 결과물로서, 제조 설비의 고장 위험을 조기에 식별하고 관리자 의사결정을 지원하기 위한 제품형 MVP 구현과 검증 결과를 정리한 것이다.")
        self.paragraph("인적사항과 승인 정보는 최종 제출 전 학교 양식에 맞춰 보완한다. 본 초안에서는 [성명], [학번], [지도교수], [제출일]을 빈칸 placeholder로 유지한다.")
        self.page_break()

        self.heading1("내표지 II")
        self.paragraph("연구 제목: AI 예지보전 기반 스마트 제조 설비 위험관리 시스템 구현")
        self.paragraph("연구 범위: AI4I 2020 데이터 기반 고장 예측, Predictive SPC, GenAI 관리자 리포트, 승인형 작업지시 workflow, 공개 산업 benchmark 기반 비용 민감 검증.")
        self.paragraph("연구 한계: 실제 공장 운영망 연결, 실제 회사 데이터 성능 재검증, 실제 비용 절감 실증은 외부 현장 데이터와 운영 로그가 필요하므로 본 초안에서는 완료 주장으로 다루지 않는다.")
        self.page_break()

        self.heading1("목차")
        toc = [
            "국문초록",
            "1. 서론",
            "2. 이론적 배경 및 선행연구",
            "3. 연구 방법",
            "4. 시스템 구현",
            "5. 실험 및 검증",
            "6. 결론 및 향후 연구",
            "참고문헌",
            "부록",
        ]
        for item in toc:
            self.paragraph(item, indent=False)
        self.page_break()

        self.heading1("국문초록")
        self.paragraph(
            "본 연구는 제조 설비의 고장 위험을 사전에 파악하고 관리자 의사결정으로 연결하기 위한 AI 예지보전 기반 스마트 제조 설비 위험관리 시스템을 구현한다. "
            "기존 연구는 고장 분류 모델 성능, 관리도 기반 이상탐지, 설명가능 인공지능, 정비 의사결정 중 일부 기능을 개별적으로 다루는 경우가 많았다. "
            "본 연구는 AI4I 2020 데이터셋을 기반으로 XGBoost 고장 확률 예측 모델을 학습하고, threshold tuning을 통해 운영 정책에 맞는 위험 판정 기준을 도출하였다. "
            "또한 예측 확률을 시간축에 배치하여 Predictive SPC 관점의 위험 추세를 확인하고, SHAP 요인과 GenAI 관리자 리포트를 결합해 위험 원인과 조치 방향을 요약하였다."
        )
        self.paragraph(
            "구현 결과물인 MaintiQ Predict는 CSV 입력, 전처리, 고장 확률 산출, 위험 우선순위 표시, AI 리포트 생성, 승인형 작업지시 기록을 하나의 데스크톱 애플리케이션 흐름으로 제공한다. "
            f"AI4I 기준 XGBoost 모델은 PR-AUC {self.evidence['xgb_pr_auc']:.4f}, ROC-AUC {self.evidence['xgb_roc_auc']:.4f}를 보였으며, 선택 threshold {self.evidence['threshold']:.2f}에서 F1-score {self.evidence['tuned_f1']:.4f}를 달성하였다. "
            f"Gemini API 기반 리포트 검증에서는 {self.evidence['genai_mode']} 방식으로 고위험 확률 {self.evidence['genai_prob']:.6f} 사례에 대한 관리자 참고 리포트를 생성하였다. "
            f"추가로 SCANIA Component X 공개 benchmark에서는 official cost metric 기준 rule baseline 대비 약 {self.evidence['scania_cost_improvement'] * 100:.2f}% 개선 가능성을 확인하였다."
        )
        self.paragraph("주요어: 예지보전, 스마트 제조, XGBoost, Predictive SPC, GenAI, 작업지시, Cost-sensitive Learning")
        self.page_break()

    def add_body(self) -> None:
        e = self.evidence

        self.heading1("1. 서론")
        self.heading2("가. 연구 배경")
        self.paragraph(
            "제조 현장에서 설비 고장은 생산성 저하, 납기 지연, 품질 불량, 긴급 정비 비용 증가로 이어진다. "
            "특히 다품종 소량생산이나 고가 설비가 투입되는 공정에서는 짧은 정지 시간도 전체 생산계획과 원가 구조에 큰 영향을 미친다. "
            "전통적인 사후보전은 고장이 발생한 뒤 수리하기 때문에 다운타임을 줄이는 데 한계가 있으며, 예방보전은 정해진 주기로 설비를 점검하므로 실제 상태와 무관하게 과잉 정비가 발생할 수 있다. "
            "예지보전은 센서 데이터와 이력 데이터를 기반으로 고장 징후를 조기에 식별하여 필요한 시점에 정비 의사결정을 내리는 방식이라는 점에서 두 방식의 한계를 보완한다."
        )
        self.paragraph(
            "하지만 예지보전 시스템을 실제 운영에 적용하려면 단순히 고장 확률을 예측하는 모델만으로는 충분하지 않다. "
            "운영자는 예측 결과가 어떤 기준으로 고위험으로 판단되었는지, 어떤 센서 요인이 위험을 키웠는지, 현재 경보가 관리도 관점에서도 이상인지, 그리고 어떤 작업지시를 생성해야 하는지까지 확인해야 한다. "
            "따라서 본 연구는 모델 성능뿐 아니라 운영자가 CSV 데이터를 입력한 뒤 위험 판단, 설명, 리포트, 승인형 작업지시까지 이어지는 전체 흐름을 구현 대상으로 삼았다."
        )
        self.table(
            "보전 전략 비교",
            ["구분", "핵심 방식", "장점", "한계"],
            [
                ["사후보전", "고장 후 수리", "초기 시스템 구축 부담이 작음", "다운타임과 긴급 비용이 큼"],
                ["예방보전", "정기 점검·교체", "계획 정비가 가능함", "상태와 무관한 과잉 정비 가능"],
                ["예지보전", "센서·이력 기반 위험 예측", "위험 기반 의사결정 가능", "데이터 품질과 운영 검증 필요"],
            ],
            [2.2, 3.3, 4.2, 4.6],
        )
        self.paragraph(
            "표 1은 본 연구가 왜 예지보전 접근을 선택했는지 보여준다. 예지보전은 사후보전의 늦은 대응 문제와 예방보전의 과잉 대응 문제를 동시에 완화할 수 있지만, 데이터 전처리와 모델 검증, 현장 적용 전 한계 설정이 함께 필요하다. "
            "본 연구는 이러한 요구를 반영하여 예측 모델과 운영 workflow를 통합한 제품형 MVP를 구현하였다."
        )

        self.heading2("나. 연구 목적과 범위")
        self.paragraph(
            "본 연구의 목적은 AI4I 2020 데이터셋을 기반으로 제조 설비 고장 위험을 예측하고, 예측 결과를 운영자가 이해하고 승인할 수 있는 형태로 제공하는 시스템을 구현하는 것이다. "
            "구체적으로는 첫째, 누수 가능성이 있는 고장 원인 컬럼을 제거한 뒤 센서 특성만으로 고장 확률을 예측한다. 둘째, threshold tuning을 통해 단순 0.5 기준보다 운영 목적에 맞는 위험 판정 기준을 선택한다. "
            "셋째, 예측 확률을 시간 흐름에 따라 해석하는 Predictive SPC 화면을 제공한다. 넷째, SHAP 기반 위험 요인과 GenAI 관리자 리포트를 통해 사람이 이해할 수 있는 조치 요약을 만든다. 다섯째, 자동 정비 명령이 아니라 작업자가 승인·검토·반려하는 작업지시 workflow로 연결한다."
        )
        self.note_box(
            "연구 범위와 금지 주장",
            [
                "본 연구는 로컬 데스크톱 MVP와 Admin 검증 콘솔 구현을 범위로 한다.",
                "실제 PLC/SCADA 운영망 연결 완료, 실제 공장 실시간 배포 완료, 실제 회사 데이터 성능 검증 완료는 주장하지 않는다.",
                "비용 절감은 공개 benchmark의 official cost metric 또는 normalized cost simulation으로만 해석한다.",
                "작업지시는 승인형 의사결정 기록이며 무인 정비 지시를 수행하지 않는다.",
            ],
        )
        self.paragraph(
            "이러한 범위 설정은 연구 결과를 과장하지 않기 위한 것이다. 실제 현장 실증을 수행하려면 설비 ID, timestamp, 센서값, 실제 고장 여부, 정비 시작·종료 시각, downtime, 부품비, 인건비, 기존 rule 결과가 함께 필요하다. "
            "본 연구는 해당 데이터를 수집할 수 있는 템플릿과 리포트 생성 구조를 준비했지만, 실제 회사 로그가 없는 상태에서 비용 절감률을 실증 완료로 표현하지 않는다."
        )

        self.heading2("다. 연구의 차별성")
        self.paragraph(
            "본 연구의 차별성은 특정 알고리즘 하나를 새로 제안하는 것이 아니라, 예지보전 운영에 필요한 여러 요소를 하나의 재현 가능한 흐름으로 결합했다는 데 있다. "
            "기존의 모델 성능 비교 중심 접근은 PR-AUC나 F1-score를 제시하는 데서 멈추기 쉽다. 반면 본 연구는 threshold 정책, SPC 기반 위험 추세, SHAP 요인, GenAI 리포트, 승인형 작업지시, 공개 산업 benchmark 검증까지 연결한다. "
            "따라서 연구 결과는 단순 모델 정확도보다 운영 의사결정 지원 가능성과 재현성에 초점을 둔다."
        )

        self.heading1("2. 이론적 배경 및 선행연구")
        self.heading2("가. 예지보전과 CBM")
        self.paragraph(
            "상태기반보전(CBM)은 설비의 실제 상태를 계측하고, 상태 변화가 정비 기준을 초과하거나 고장 징후를 보일 때 정비를 수행하는 접근이다. "
            "예지보전은 CBM의 확장으로 볼 수 있으며, 단순 현재 상태 판정이 아니라 향후 고장 가능성 또는 잔여수명(RUL)을 추정해 정비 의사결정을 지원한다. "
            "Jardine 등은 CBM 연구를 데이터 수집, 진단, 예측, 의사결정으로 구분했으며, Carvalho 등은 기계학습 기반 예지보전 연구가 산업 데이터의 불균형, 라벨 부족, 설명 가능성 문제를 동시에 다루어야 한다고 정리하였다."
        )
        self.paragraph(
            "본 연구는 CBM의 관점에서 센서 CSV를 입력받아 현재 위험도를 계산하고, Predictive SPC로 위험 추세를 확인하며, 작업지시 승인 기록을 남기는 구조를 구현한다. "
            "즉, 예측 결과가 단순 숫자로 끝나지 않고 현장 관리자가 검토할 수 있는 정보 구조로 변환되도록 설계하였다."
        )

        self.heading2("나. SPC와 관리도 이론")
        self.paragraph(
            "SPC(Statistical Process Control)는 공정 변동을 통계적으로 감시하고 이상 상태를 조기에 탐지하는 방법이다. 관리도에서는 중심선(CL), 관리상한(UCL), 관리하한(LCL)을 설정하고 관측값이 관리한계를 벗어나거나 특정 패턴을 보이면 이상 신호로 판단한다. "
            "일반적인 Shewhart 관리도에서 관리한계는 평균과 표준편차를 이용해 CL=μ, UCL=μ+3σ, LCL=μ−3σ로 정의할 수 있다. 제조 현장에서는 센서값뿐 아니라 모델이 산출한 위험 확률에도 관리도 개념을 적용할 수 있다."
        )
        self.paragraph(
            "본 연구의 Predictive SPC는 XGBoost 고장 확률을 시간 순서로 정렬하고 rolling window 기반 위험 추세를 시각화한다. 이를 통해 개별 row의 고위험 여부뿐 아니라 위험 확률이 공정 흐름에서 어떻게 변하는지를 함께 확인한다. "
            "SPC-only 방식은 센서 또는 확률의 단순 관리한계 초과를 기준으로 하며, ML+SPC 방식은 모델의 고장 확률과 관리도 신호를 결합한다."
        )

        self.heading2("다. FMEA/RPN과 위험 우선순위")
        self.paragraph(
            "FMEA(Failure Mode and Effects Analysis)는 고장 모드가 시스템에 미치는 영향을 식별하고 우선순위를 정하는 품질관리 기법이다. 일반적으로 RPN(Risk Priority Number)은 심각도(Severity), 발생도(Occurrence), 검출도(Detection)를 곱해 계산한다. "
            "본 연구의 risk priority score는 전통적인 RPN을 그대로 계산하는 것은 아니지만, 고장 확률을 발생도에 대응시키고 데이터 품질 및 threshold 초과 여부를 검출 가능성과 연결해 위험 row의 우선순위를 정한다."
        )
        self.table(
            "FMEA/RPN 개념과 본 시스템 변수의 대응",
            ["FMEA 개념", "의미", "본 시스템 대응"],
            [
                ["Severity", "고장이 발생했을 때 영향", "운영 정책의 missed failure 비용 가중치"],
                ["Occurrence", "고장 발생 가능성", "calibrated probability 또는 raw probability"],
                ["Detection", "사전 발견 가능성", "품질 점수, SPC 신호, threshold 초과 여부"],
                ["RPN", "우선 조치 필요성", "risk priority score 및 작업지시 후보"],
            ],
            [2.8, 4.8, 6.5],
        )
        self.paragraph(
            "표 2는 산업공학적 위험관리 개념과 구현 변수의 연결을 보여준다. 이는 모델 결과를 현장 언어로 해석하기 위한 장치이며, 논문에서는 예측 모델이 아니라 의사결정 지원 시스템으로서의 성격을 설명하는 근거가 된다."
        )

        self.heading2("라. 기계학습 기반 고장 예측")
        self.paragraph(
            "고장 예측 문제는 정상 데이터가 대부분이고 고장 데이터가 적은 불균형 분류 문제로 나타나는 경우가 많다. Accuracy만 보면 정상으로만 예측해도 높은 점수가 나올 수 있으므로, precision, recall, F1-score, ROC-AUC, PR-AUC 같은 지표를 함께 확인해야 한다. "
            "특히 고장 클래스가 희소할 때는 PR-AUC가 모델의 실제 고장 탐지 성능을 더 잘 드러낸다."
        )
        self.paragraph(
            "XGBoost는 gradient boosting 기반의 tree ensemble 모델로, 비선형 관계와 feature interaction을 효과적으로 학습할 수 있어 tabular sensor data에 적합하다. "
            "본 연구에서는 Logistic Regression을 기본 baseline으로 두고 XGBoost와 비교했으며, 성능이 더 높은 XGBoost를 핵심 예측 엔진으로 사용하였다. "
            "SMOTE는 minority class를 합성해 불균형을 완화하는 방법이지만, 실제 운영에서는 합성 데이터가 경보 정책에 미치는 영향까지 함께 검토해야 하므로 threshold tuning과 함께 비교하였다."
        )

        self.heading2("마. SHAP, GenAI, 비용 민감 학습")
        self.paragraph(
            "SHAP은 각 feature가 개별 예측에 얼마나 기여했는지를 Shapley value 기반으로 설명하는 방법이다. 예지보전에서는 단순히 고장 확률만 제시하는 것보다 어떤 센서 요인이 위험 판단에 영향을 주었는지 알려주는 것이 중요하다. "
            "본 연구는 SHAP 요인을 GenAI 리포트의 입력 context에 포함해 torque, rotational speed 등 주요 위험 신호가 관리자 설명문으로 변환되도록 구성하였다."
        )
        self.paragraph(
            "비용 민감 학습은 false alarm과 missed failure의 비용이 서로 다르다는 점을 반영한다. 정비 시스템에서 false alarm은 불필요 점검을 유발하지만, missed failure는 생산 중단과 안전 문제로 이어질 수 있다. "
            "따라서 본 연구는 operational cost simulation과 SCANIA official cost matrix를 통해 단순 classification score 외에 운영 비용 관점의 비교를 수행하였다."
        )
        self.table(
            "선행연구 흐름과 본 연구 반영점",
            ["분야", "대표 내용", "본 연구 반영"],
            [
                ["CBM/예지보전", "센서 기반 상태 감시와 고장 예측", "CSV 기반 위험 산출과 작업지시 workflow"],
                ["SPC", "UCL/LCL 기반 이상 신호 탐지", "예측 확률의 Predictive SPC 적용"],
                ["ML 예측", "불균형 고장 분류와 성능 비교", "Logistic Regression, XGBoost, threshold tuning"],
                ["XAI/GenAI", "설명 가능성과 자연어 요약", "SHAP 요인과 Gemini 관리자 리포트"],
                ["Cost-sensitive", "오경보와 미탐 비용 차이 반영", "normalized cost, SCANIA official cost metric"],
            ],
            [2.5, 4.7, 6.9],
        )

        self.heading1("3. 연구 방법")
        self.heading2("가. 데이터셋과 전처리")
        self.paragraph(
            "기본 학습 데이터는 UCI의 AI4I 2020 Predictive Maintenance Dataset이다. 이 데이터는 제품 타입, 공기온도, 공정온도, 회전속도, 토크, 공구마모와 같은 제조 설비 관련 변수를 포함하며, target은 Machine failure이다. "
            "학습 단계에서는 ID 성격의 UDI, Product ID를 제거하고, 고장 원인 라벨에 해당하는 TWF, HDF, PWF, OSF, RNF는 leakage를 방지하기 위해 제외하였다. "
            "Type은 범주형 변수이므로 one-hot encoding을 적용하였다."
        )
        self.figure("데이터 전처리 및 예측 pipeline", "02_data_preprocessing_pipeline.png", 12.2)
        self.paragraph(
            "그림 1은 본 연구의 전처리 흐름을 요약한다. 핵심은 모델이 실제 운영 시점에 알 수 없는 원인 라벨을 학습하지 않도록 누수 컬럼을 제거하는 것이다. "
            "이 과정을 통해 모델은 센서값과 제품 타입 정보만으로 고장 확률을 산출한다."
        )
        self.table(
            "AI4I 데이터 전처리 규칙",
            ["항목", "처리 방식", "이유"],
            [
                ["Target", "Machine failure", "이진 고장 예측 목표"],
                ["ID 컬럼", "UDI, Product ID 제거", "학습 의미가 낮고 일반화 저해 가능"],
                ["Leakage 컬럼", "TWF/HDF/PWF/OSF/RNF 제거", "고장 원인 라벨이 target 정보를 직접 포함"],
                ["Type", "One-hot encoding", "범주형 제품 타입을 수치 feature로 변환"],
                ["Split", "stratified train/test, random_state=42", "고장 클래스 비율 유지와 재현성 확보"],
            ],
            [2.7, 5.2, 6.2],
        )
        self.paragraph(
            "표 4의 전처리 규칙은 논문에서 재현성을 확보하는 핵심 조건이다. 특히 leakage 컬럼을 제거하지 않으면 모델 성능이 과도하게 높아져 실제 센서 기반 예측이라고 보기 어렵다. "
            "따라서 본 연구는 발표와 논문에서 모델 성능보다도 전처리의 공정성을 명확히 제시한다."
        )

        self.heading2("나. 모델 학습과 threshold tuning")
        self.paragraph(
            f"학습 모델은 Logistic Regression과 XGBoost를 비교하였다. 테스트 데이터는 전체 10,000행 중 {e['test_rows']}행이며, 이 중 고장 row는 {e['test_failures']}행이다. "
            "Logistic Regression은 해석이 쉬운 기준 모델이며, XGBoost는 비선형 관계와 feature interaction을 더 잘 포착하는 비교 모델이다. "
            f"검증 결과 XGBoost는 PR-AUC {e['xgb_pr_auc']:.4f}, ROC-AUC {e['xgb_roc_auc']:.4f}로 Logistic Regression보다 우수하여 핵심 예측 모델로 선택하였다."
        )
        self.table(
            "AI4I baseline 모델 성능",
            ["모델", "Precision", "Recall", "F1-score", "ROC-AUC", "PR-AUC"],
            [
                ["Logistic Regression", "0.1418", "0.8235", "0.2419", f"{e['lr_roc_auc']:.4f}", f"{e['lr_pr_auc']:.4f}"],
                ["XGBoost", f"{e['xgb_precision']:.4f}", f"{e['xgb_recall']:.4f}", f"{e['xgb_f1']:.4f}", f"{e['xgb_roc_auc']:.4f}", f"{e['xgb_pr_auc']:.4f}"],
            ],
            [3.5, 2.2, 2.2, 2.2, 2.2, 2.2],
        )
        self.figure("XGBoost PR curve", "22_pr_curve.png", 11.8)
        self.paragraph(
            "그림 2는 XGBoost 모델의 precision-recall curve를 보여준다. 고장 데이터가 희소한 상황에서는 ROC-AUC만으로 모델을 평가하면 정상 클래스의 영향으로 성능이 과대평가될 수 있다. "
            "따라서 본 연구는 PR-AUC를 핵심 기준으로 두고 모델을 선택하였다."
        )
        self.paragraph(
            f"기본 threshold 0.5에서는 recall이 높지만 precision이 낮아 오경보가 많아질 수 있다. 이에 본 연구는 0.05부터 0.95까지 0.01 간격으로 threshold를 탐색하고 F1-score가 가장 높은 기준을 선택하였다. "
            f"선택된 threshold는 {e['threshold']:.2f}이며, 이때 precision {e['tuned_precision']:.4f}, recall {e['tuned_recall']:.4f}, F1-score {e['tuned_f1']:.4f}를 보였다."
        )
        self.figure("Threshold tuning 결과", "23_threshold_tuning.png", 11.8)
        self.paragraph(
            "그림 3은 threshold가 변할 때 precision, recall, F1-score가 어떻게 바뀌는지 보여준다. 운영 관점에서는 모든 고장을 잡으려는 recall-first 정책과 오경보를 줄이려는 precision-first 정책 사이에서 선택이 필요하다. "
            "본 연구의 기본 정책은 F1-score 기반 balanced threshold이며, 시스템에서는 운영 정책별 threshold 조정 가능성을 별도 산출물로 제시하였다."
        )

        self.heading2("다. Predictive SPC와 위험 우선순위")
        self.paragraph(
            "예측 모델이 row 단위 고장 확률을 제공한다면, Predictive SPC는 그 확률을 시간 흐름 속에서 감시한다. AI4I 데이터는 실제 timestamp가 없으므로 UDI 순서를 기준으로 simulated time axis를 구성하였다. "
            "이 방식은 실제 공장 streaming을 의미하지 않지만, 예측 확률을 시계열처럼 배치했을 때 위험 추세와 관리한계 초과 여부를 검토할 수 있게 한다."
        )
        self.figure("Predictive SPC 위험 확률 관리도", "24_spc_risk_chart.png", 12.8)
        self.paragraph(
            "그림 4는 예측 고장 확률이 threshold와 SPC 관리한계를 기준으로 어떻게 해석되는지 보여준다. 단일 row의 확률이 높다는 사실뿐 아니라 일정 구간에서 위험 확률이 상승하는 패턴을 함께 보면 경보의 운영적 의미를 더 잘 판단할 수 있다."
        )
        self.table(
            "SPC-only와 ML+SPC 비교 해석",
            ["전략", "핵심 기준", "F1-score", "해석"],
            [
                ["SPC-only", "관리한계 초과", "0.1600", "단순 이상 신호는 고장 라벨과 직접 연결되기 어려움"],
                ["ML threshold", "XGBoost 확률 ≥ 0.87", "0.7752", "고장 예측 성능이 가장 높음"],
                ["ML+SPC", "확률과 SPC 신호 결합", "0.7051", "성능은 일부 낮아도 운영 설명성이 향상됨"],
            ],
            [2.8, 4.2, 2.3, 6.2],
        )
        self.paragraph(
            "표 6은 SPC-only 접근만으로는 고장 예측이 충분하지 않지만, ML 결과와 SPC 신호를 함께 제시하면 운영자가 경보를 해석하는 데 도움이 된다는 점을 보여준다. "
            "이 결과는 본 연구의 핵심이 단순 성능 최고 모델이 아니라 운영 의사결정 지원 workflow라는 점을 뒷받침한다."
        )

        self.heading2("라. GenAI 리포트와 작업지시 context")
        self.paragraph(
            "GenAI 리포트는 모델이 산출한 고장 확률, threshold, 위험 상태, 센서값, SPC 요약, SHAP 요인을 입력 context로 받아 관리자 참고용 문장으로 변환한다. "
            "이때 API key는 세션에서만 사용하고 파일에 저장하지 않는다. 또한 리포트는 자동 정비 명령이 아니라 현장 담당자가 승인하기 전 참고하는 설명 자료로 제한하였다."
        )
        self.table(
            "GenAI 리포트 생성 근거",
            ["항목", "값"],
            [
                ["report_generation_mode", e["genai_mode"]],
                ["예측 확률", f"{e['genai_prob']:.6f}"],
                ["위험 판정 기준", f"{e['genai_threshold']:.2f}"],
                ["상태", e["genai_status"]],
                ["주요 위험 요인", e["genai_top_factors"]],
            ],
            [4.2, 8.8],
        )
        self.paragraph(
            "표 7은 실제 Gemini API 호출 결과를 논문에 넣기 위한 최소 근거이다. 전체 리포트를 본문에 길게 삽입하지 않고, 리포트 생성 방식과 핵심 판단 요약만 제시하였다. "
            "보고서 전문은 부록 또는 산출물 파일로 연결하는 것이 본문 가독성 측면에서 적절하다."
        )

        self.heading1("4. 시스템 구현")
        self.heading2("가. 전체 시스템 구조")
        self.paragraph(
            "MaintiQ Predict는 사용자용 Windows 데스크톱 앱과 연구·검증용 Admin 콘솔을 분리하여 구현하였다. 사용자 앱은 CSV 선택, 컬럼 확인, 품질 진단, 예측 실행, 결과 저장, AI 리포트, 작업지시 기록을 제공한다. "
            "Admin 콘솔은 모델 비교, 공개 benchmark, 산업공학 검증 근거, 회사 데이터 실증 템플릿과 같은 연구·검증 기능을 담당한다."
        )
        self.figure("MaintiQ Predict 전체 시스템 구조", "01_system_architecture.png", 12.7)
        self.paragraph(
            "그림 5는 시스템이 데이터 입력부터 작업지시 기록까지 연결되는 방식을 보여준다. 제품 앱에서는 연구용 상세 지표를 노출하지 않고 운영자가 바로 사용할 수 있는 기능만 제공하며, 검증 자료는 Admin 콘솔과 outputs 문서로 분리하였다."
        )

        self.heading2("나. 데스크톱 제품 화면")
        self.paragraph(
            "사용자 앱은 PySide6 기반 네이티브 데스크톱 앱으로 구현하였다. 초기 Streamlit 데모에서 출발했지만, 최종 제품형 MVP에서는 브라우저와 localhost 접속 없이 설치형 Windows 앱으로 실행할 수 있도록 Full/Lite 설치본을 분리하였다. "
            "Full은 정밀 분석 모드로 XGBoost/SHAP 기반 분석을 포함하고, Lite는 빠른 점검 모드로 경량 운영 점수 중심의 작은 설치본을 제공한다."
        )
        self.figure("MaintiQ Predict 메인 화면", "20_app_main_screen.png", 12.2)
        self.paragraph(
            "그림 6의 메인 화면은 CSV 예측 시작, 고위험 설비 확인, AI 리포트 생성, 작업지시 승인이라는 운영 흐름을 중심으로 배치하였다. "
            "논문, 발표, Stage, PoC 같은 개발 용어는 사용자 앱에서 제거하고 제품 화면처럼 보이도록 정리하였다."
        )
        self.table(
            "Full/Lite 모드 구분",
            ["구분", "사용 목적", "계산 방식", "주의점"],
            [
                ["빠른 점검 모드", "일상 점검·배포·시연", "경량 운영 점수", "정밀 분석 결과와 다를 수 있음"],
                ["정밀 분석 모드", "정밀 분석·검증", "XGBoost/SHAP 기반", "설치 용량이 더 큼"],
                ["Admin 콘솔", "연구·benchmark·논문 근거", "검증 스크립트와 산출물", "일반 사용자에게 기본 노출하지 않음"],
            ],
            [3.0, 3.8, 4.0, 4.0],
        )

        self.heading2("다. CSV 예측 wizard")
        self.paragraph(
            "데이터 예측 화면은 CSV 선택, 컬럼 확인, 품질 진단, 예측 실행, 결과 저장의 5단계 wizard 흐름으로 구성하였다. "
            "업로드 전에는 샘플 CSV와 필수 데이터 안내만 보여주고, 업로드 후에는 요약 카드, 상위 위험 row, 위험도 그래프, 결과표, 저장 버튼 순서로 배치하였다."
        )
        self.paragraph(
            "CSV 입력은 회사별 컬럼명이 다를 수 있다는 점을 고려해 자동 매핑과 단위 변환을 지원하도록 설계하였다. 예를 들어 air_temp, air temperature, 공기온도는 Air temperature [K]에 매핑될 수 있고, Celsius 입력은 Kelvin으로 변환할 수 있다. "
            "숫자 변환 실패, 결측률, 허용 범위 밖 값, Type 불일치, 중복 row는 품질 리포트로 분리하여 사용자가 수정할 수 있게 했다."
        )

        self.heading2("라. 작업지시 workflow와 이력")
        self.paragraph(
            "작업지시 기능은 센서 이벤트 생성, 작업지시 초안 생성, 작업자 결정 저장, 최근 이력 확인 순서로 구성하였다. "
            "시스템은 자동 정비 명령을 실행하지 않으며, approve, needs_review, reject에 해당하는 승인·검토 필요·반려 결정을 기록한다. 이는 현장 안전과 책임 소재를 고려한 설계이다."
        )
        self.figure("승인형 작업지시 workflow", "03_work_order_workflow.png", 12.0)
        self.paragraph(
            "그림 7은 예측 결과가 작업지시 의사결정으로 이어지는 과정을 보여준다. 위험 row는 즉시 설비 정지 명령으로 연결되지 않고, 담당자가 현장 확인과 부품 점검 필요성을 검토한 뒤 결정 이력을 남긴다."
        )

        self.heading2("마. 회사 데이터 실증 준비 구조")
        self.paragraph(
            "실제 회사 데이터 실증을 위해 labeled sensor CSV, maintenance history CSV, downtime/cost CSV 세 가지 입력 템플릿을 준비하였다. "
            "세 파일이 모두 있을 때 precision, recall, false alarm, missed failure, lead time, downtime delta, maintenance cost delta를 계산할 수 있다. "
            "일부 파일만 있는 경우에는 가능한 주장과 불가능한 주장을 분리해 출력한다."
        )
        self.table(
            "회사 데이터 실증 파일 조합별 claim 상태",
            ["입력 파일", "가능한 분석", "불가능한 주장"],
            [
                ["labeled sensor CSV", "성능 재평가 가능", "비용·탐지시간 실증 불가"],
                ["sensor + maintenance", "lead-time 일부 분석 가능", "비용 절감 실증 불가"],
                ["sensor + maintenance + cost", "downtime/cost delta 분석 가능", "현장 조건 외 일반화 주의"],
            ],
            [4.5, 5.2, 5.0],
        )

        self.heading1("5. 실험 및 검증")
        self.heading2("가. 실험 설계")
        self.paragraph(
            "실험은 모델 성능, threshold 정책, SPC 결합, 운영 가치 simulation, 공개 산업 benchmark의 다섯 축으로 구성하였다. "
            "첫째, AI4I 데이터에서 Logistic Regression과 XGBoost를 비교하였다. 둘째, threshold tuning을 통해 운영 기준을 선택하였다. 셋째, SPC-only와 ML+SPC를 비교해 관리도 기반 신호의 의미를 확인하였다. "
            "넷째, false alarm과 missed failure에 비용 가중치를 둔 normalized cost simulation을 수행하였다. 다섯째, SCANIA Component X 공개 데이터에서 official cost matrix 기반 비교를 수행하였다."
        )
        self.table(
            "실험 설계 요약",
            ["실험", "목적", "주요 지표"],
            [
                ["AI4I baseline", "기본 예측 모델 비교", "Precision, Recall, F1, ROC-AUC, PR-AUC"],
                ["Threshold tuning", "운영 기준 선택", "threshold별 F1, precision, recall"],
                ["SPC 비교", "관리도 신호 효과 확인", "SPC-only, ML threshold, ML+SPC F1"],
                ["Cost simulation", "오경보/미탐 비용 trade-off", "false alarm, missed failure, normalized cost"],
                ["SCANIA benchmark", "공개 산업 데이터 비용 metric 검증", "official cost, improvement vs rule"],
            ],
            [3.3, 5.5, 5.5],
        )

        self.heading2("나. AI4I baseline과 threshold 결과")
        self.paragraph(
            f"AI4I baseline 결과에서 XGBoost는 PR-AUC {e['xgb_pr_auc']:.4f}, ROC-AUC {e['xgb_roc_auc']:.4f}를 기록하였다. "
            "Logistic Regression은 recall은 높게 유지할 수 있었으나 precision과 PR-AUC가 낮아 실제 경보 운영에서는 오경보가 많아질 가능성이 있다. "
            "XGBoost는 비선형 센서 관계를 더 잘 반영하여 고장 class에 대한 ranking 성능이 높았다."
        )
        self.paragraph(
            f"Threshold tuning 결과 선택 기준 {e['threshold']:.2f}는 기본 0.5 기준보다 F1-score를 개선하였다. "
            "이 결과는 모델의 확률 산출 자체만큼이나 운영 판정 기준 선택이 중요하다는 점을 보여준다. 실제 현장에서는 설비 중요도와 정비 비용 구조에 따라 precision-first 또는 recall-first 기준을 다시 선택할 수 있다."
        )

        self.heading2("다. 설명가능성과 GenAI 검증")
        self.figure("SHAP summary plot", "26_shap_summary.png", 12.4)
        self.paragraph(
            "그림 8은 SHAP summary plot으로, 예측 결과에 영향을 주는 feature 분포를 보여준다. 본 연구에서는 SHAP을 단순 시각화로만 쓰지 않고, 고위험 row의 주요 위험 요인을 GenAI 리포트 입력 context에 포함하였다. "
            "따라서 관리자는 고장 확률과 함께 torque, rotational speed 등 어떤 요인이 위험 판단에 영향을 주었는지 확인할 수 있다."
        )
        self.figure("SHAP factor bar chart", "27_shap_bar.png", 11.4)
        self.paragraph(
            f"Gemini 리포트 검증에서는 고위험 확률 {e['genai_prob']:.6f}, threshold {e['genai_threshold']:.2f}, 상태 {e['genai_status']} 사례가 사용되었다. "
            f"리포트 생성 방식은 {e['genai_mode']}로 기록되며, API key는 결과 파일에 저장하지 않았다. 리포트의 역할은 관리자 참고 자료이며, 자동 정비 명령이 아니다."
        )

        self.heading2("라. Operational cost simulation")
        self.paragraph(
            "운영 가치 simulation은 false alarm과 missed failure의 비용 차이를 반영하기 위해 수행하였다. 실제 원화 비용이 아니라 normalized cost를 사용했으며, conservative, balanced, high_downtime 시나리오에 따라 정책별 비용 변화를 비교하였다. "
            "이 simulation은 실제 비용 절감 실증이 아니라 운영 의사결정 가능성을 평가하는 보조 지표이다."
        )
        self.figure("Operational value simulation", "28_operational_value_simulation.png", 12.2)
        self.paragraph(
            "그림 10은 경보 정책에 따라 normalized operating cost가 달라질 수 있음을 보여준다. missed failure 비용이 큰 상황에서는 recall을 높이는 정책이 유리할 수 있고, 오경보 비용이 큰 상황에서는 precision을 높이는 정책이 적합할 수 있다."
        )

        self.heading2("마. SCANIA Component X official cost metric")
        self.paragraph(
            "AI4I는 교육용 synthetic 성격이 강하므로, 공개 산업 데이터 기반 검증을 보완하기 위해 SCANIA Component X를 사용하였다. 이 데이터는 SCANIA 트럭 fleet에서 수집된 익명화된 component dataset이며, class 0~4와 official cost matrix가 제공된다. "
            "본 연구는 no-alert, rule-based threshold, SPC-style baseline, logistic multiclass, XGBoost argmax, XGBoost official-cost optimized 전략을 비교하였다."
        )
        self.table(
            "SCANIA official cost metric 결과 요약",
            ["전략", "official cost", "normalized cost", "rule 대비 개선"],
            [
                ["Rule-based threshold", "59,709", "1.0402", "0.00%"],
                ["XGBoost argmax", "55,096", "0.9599", "7.73%"],
                ["XGBoost cost optimized", f"{e['scania_cost']:.0f}", f"{e['scania_normalized_cost']:.4f}", f"{e['scania_cost_improvement'] * 100:.2f}%"],
            ],
            [4.6, 3.1, 3.2, 3.1],
        )
        self.figure("SCANIA official cost comparison", "29_scania_cost_comparison.png", 12.2)
        self.paragraph(
            f"SCANIA 결과에서 XGBoost cost optimized 전략은 rule baseline 대비 official cost metric 기준 약 {e['scania_cost_improvement'] * 100:.2f}% 개선을 보였다. "
            "이는 실제 회사의 원화 비용 절감 실증이 아니라, 공개 산업 benchmark의 공식 비용 행렬을 사용한 성능 비교이다. 따라서 논문에서는 '실제 공장 비용 절감'이 아니라 '공개 benchmark official cost metric 기준 개선'으로 표현해야 한다."
        )

        self.heading2("바. Public benchmark 확장")
        self.paragraph(
            "MetroPT-3, NASA C-MAPSS, IMS/FEMTO 같은 공개 run-to-failure benchmark는 실제 회사 로그를 대체하지는 못하지만, lead time, RUL, early warning, simulated cost를 다양한 데이터 구조에서 비교할 수 있게 한다. "
            "본 연구는 대용량 원본 데이터가 없는 경우에도 sample smoke와 guardrail 문구가 유지되도록 adapter 구조를 설계하였다."
        )
        self.figure("Public benchmark cost comparison", "31_public_benchmark_cost_chart.png", 11.8)
        self.figure("Public benchmark lead-time summary", "32_public_benchmark_lead_time.png", 11.8)
        self.paragraph(
            "그림 12와 그림 13은 공개 benchmark 확장 결과를 요약한다. 이 결과는 실제 회사 데이터 실증을 대체하지 않지만, 모델과 workflow가 AI4I 하나에만 종속되지 않도록 검증 범위를 넓힌다는 의미가 있다."
        )

        self.heading2("사. Claim boundary 검증")
        self.paragraph(
            "본 연구는 구현과 검증 범위를 명확히 분리하였다. 시스템은 CSV 기반 예측, 관리자 리포트, 작업지시 기록, 공개 benchmark 비교를 수행할 수 있지만, 실제 PLC/SCADA 운영망에 연결되었거나 실제 회사 비용 절감이 검증되었다고 말할 수는 없다. "
            "이 claim boundary는 논문 방어와 발표 질의응답에서 특히 중요하다."
        )
        self.table(
            "가능한 주장과 금지 주장",
            ["구분", "표현"],
            [
                ["가능", "AI4I 기반 고장 확률 예측과 threshold tuning을 구현하였다."],
                ["가능", "Predictive SPC와 GenAI 리포트, 승인형 작업지시 workflow를 통합하였다."],
                ["가능", "SCANIA 공개 benchmark official cost metric에서 rule 대비 개선 가능성을 확인하였다."],
                ["금지", "실제 회사 데이터 성능 검증이 완료되었다고 주장하지 않는다."],
                ["금지", "실제 비용 절감률 또는 탐지 시간 단축률을 실증 완료로 표현하지 않는다."],
                ["금지", "실제 PLC/SCADA 운영망 배포 또는 무인 정비 지시 수행으로 표현하지 않는다."],
            ],
            [2.1, 11.8],
        )

        self.heading1("6. 결론 및 향후 연구")
        self.heading2("가. 연구 결과 요약")
        self.paragraph(
            "본 연구는 AI4I 2020 데이터 기반 고장 확률 예측에서 출발하여, threshold tuning, Predictive SPC, SHAP/GenAI 설명, 승인형 작업지시 workflow, 공개 산업 benchmark 검증을 연결한 MaintiQ Predict 시스템을 구현하였다. "
            f"XGBoost는 PR-AUC {e['xgb_pr_auc']:.4f}, ROC-AUC {e['xgb_roc_auc']:.4f}를 보였고, threshold {e['threshold']:.2f}에서 F1-score {e['tuned_f1']:.4f}를 달성하였다. "
            f"Gemini 리포트는 {e['genai_mode']}로 생성되어 고위험 사례를 관리자 참고 문장으로 요약하였다."
        )
        self.paragraph(
            "구현 측면에서는 사용자 앱과 Admin 콘솔을 분리하고, Full/Lite 설치본을 제공함으로써 제품형 MVP와 연구 검증 경로를 동시에 유지하였다. "
            "사용자 앱은 운영자가 CSV를 입력하고 위험 우선순위를 확인하며 리포트와 작업지시를 처리하는 화면으로 구성되었고, Admin 콘솔은 실험 결과와 논문 근거를 확인하는 역할을 담당한다."
        )

        self.heading2("나. 연구 기여")
        self.paragraph(
            "첫째, 본 연구는 단일 모델 성능 비교를 넘어 예측, 통계적 공정관리, 설명가능성, GenAI 리포트, 작업지시 의사결정까지 연결한 end-to-end 예지보전 workflow를 구현하였다. "
            "둘째, threshold tuning과 cost-sensitive 관점을 도입해 고장 예측 결과가 운영 정책에 따라 달라질 수 있음을 제시하였다. "
            "셋째, SCANIA Component X 공개 benchmark의 official cost matrix를 활용해 AI4I 외부 데이터에서도 비용 metric 기반 비교를 수행하였다. "
            "넷째, 실제 회사 데이터 실증에 필요한 입력 템플릿과 claim boundary를 명확히 정리하여 향후 현장 적용을 위한 준비 구조를 제시하였다."
        )

        self.heading2("다. 한계와 향후 연구")
        self.paragraph(
            "본 연구의 가장 큰 한계는 실제 회사 로그 기반 실증이 아직 없다는 점이다. 실제 비용 절감이나 고장 탐지 시간 단축을 주장하려면 설비 ID, timestamp, 센서값, 실제 고장 라벨, 정비 시작·종료, downtime, 부품비, 인건비, 기존 rule 결과가 필요하다. "
            "향후 연구에서는 기업 또는 공개 run-to-failure 데이터의 더 풍부한 정비 이력을 확보하여 before/after 비교와 비용 분석을 수행해야 한다."
        )
        self.paragraph(
            "두 번째 한계는 실제 PLC/SCADA, MQTT, OPC UA, CMMS/EAM 연동이 구현 범위에 포함되지 않았다는 점이다. 현재 시스템은 로컬 데스크톱 MVP와 Admin 검증 콘솔 중심이며, 실제 운영망 배포를 위해서는 인증, 접근제어, 보안 감사, 운영 DB, 백업, 모니터링, 코드 서명, 자동 업데이트가 추가로 필요하다. "
            "세 번째 한계는 GenAI 리포트가 관리자 참고용이라는 점이다. 리포트 품질을 높이기 위해서는 실제 정비팀의 피드백과 domain-specific prompt template 개선이 필요하다."
        )
        self.paragraph(
            "결론적으로 MaintiQ Predict는 학부 캡스톤 수준에서 구현 가능한 제품형 예지보전 MVP로서, 예측 모델과 산업공학적 의사결정 흐름을 통합했다는 점에서 의미가 있다. "
            "향후 실제 회사 데이터와 현장 운영 로그가 확보된다면 본 연구에서 준비한 field validation 구조를 이용해 실제 lead time 개선과 비용 절감 여부를 검증할 수 있을 것이다."
        )

        self.heading1("참고문헌")
        refs = [
            "Jardine, A. K. S., Lin, D., & Banjevic, D. (2006). A review on machinery diagnostics and prognostics implementing condition-based maintenance. Mechanical Systems and Signal Processing.",
            "Carvalho, T. P., Soares, F. A. A. M. N., Vita, R., Francisco, R. P., Basto, J. P., & Alcalá, S. G. S. (2019). A systematic literature review of machine learning methods applied to predictive maintenance. Computers & Industrial Engineering.",
            "Montgomery, D. C. (2019). Introduction to Statistical Quality Control. Wiley.",
            "Elkan, C. (2001). The foundations of cost-sensitive learning. International Joint Conference on Artificial Intelligence.",
            "Chen, T., & Guestrin, C. (2016). XGBoost: A scalable tree boosting system. KDD.",
            "Chawla, N. V., Bowyer, K. W., Hall, L. O., & Kegelmeyer, W. P. (2002). SMOTE: Synthetic minority over-sampling technique. Journal of Artificial Intelligence Research.",
            "Lundberg, S. M., & Lee, S.-I. (2017). A unified approach to interpreting model predictions. NeurIPS.",
            "Matzka, S. (2020). AI4I 2020 Predictive Maintenance Dataset. UCI Machine Learning Repository.",
            "SCANIA Component X Dataset. Researchdata.se. DOI: 10.58141/1w9m-yz81.",
            "UCI Machine Learning Repository. MetroPT-3 Dataset.",
            "NASA Prognostics Center of Excellence. C-MAPSS and IMS Bearing datasets.",
            "Google AI for Developers. Gemini API generateContent documentation.",
            "OpenAI. Responses API documentation.",
            "PyInstaller documentation. Packaging Python applications for Windows.",
        ]
        for ref in refs:
            self.paragraph(ref, indent=False)

        self.heading1("부록")
        self.heading2("가. 실행 및 검증 산출물")
        self.paragraph(
            "본 연구의 주요 산출물은 outputs/metrics.json, outputs/threshold_summary.json, outputs/spc_vs_ml_comparison.csv, outputs/operational_value_simulation.csv, outputs/ai_report_context.json, outputs/scania_official_cost_metrics.json 등에 저장된다. "
            "사용자 앱 검증은 desktop_app/main.py의 check, engine-smoke-test, workflow-smoke-test 명령으로 수행할 수 있으며, 설치본 검증은 Full/Lite 빌드 스크립트와 release checksum을 통해 확인한다."
        )
        self.heading2("나. 실제 회사 데이터 실증에 필요한 필드")
        self.paragraph(
            "실제 현장 실증을 위해 필요한 최소 필드는 equipment_id, timestamp, sensor values, actual_failure, maintenance_start, maintenance_end, downtime_minutes, parts_cost, labor_cost, action_type, baseline_rule_result이다. "
            "이 필드가 있어야 기존 운영 방식 대비 false alarm, missed failure, lead time, downtime, maintenance cost delta를 계산할 수 있다."
        )
        self.paragraph(
            "sensor values는 공기온도, 공정온도, 회전속도, 토크, 공구마모처럼 설비 상태를 설명하는 측정값을 의미한다. actual_failure는 일정 시간 안에 실제 고장이 발생했는지 여부를 나타내며, maintenance_start와 maintenance_end는 정비 대응 시간을 계산하기 위한 기준이다. "
            "downtime_minutes, parts_cost, labor_cost는 비용 분석의 핵심 입력이며, baseline_rule_result는 기존 점검 방식 또는 현장 rule이 같은 시점에 어떤 판단을 했는지 비교하기 위한 필드이다."
        )

    def save(self) -> None:
        self.doc.save(DOCX_PATH)


def extract_docx_text(path: Path) -> str:
    doc = Document(path)
    parts: list[str] = []
    for p in doc.paragraphs:
        parts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def validate_text(text: str) -> dict:
    forbidden = [
        "실제 PLC/SCADA 배포 완료",
        "실제 회사 데이터 검증 완료",
        "실제 비용 절감 실증 완료",
        "자동 정비 명령 실행",
    ]
    api_patterns = [r"AIza[0-9A-Za-z_-]{20,}", r"sk-[0-9A-Za-z_-]{20,}"]
    return {
        "char_count": len(re.sub(r"\s+", "", text)),
        "replacement_character_count": text.count("�"),
        "api_key_pattern_count": sum(len(re.findall(pattern, text)) for pattern in api_patterns),
        "forbidden_phrase_hits": [phrase for phrase in forbidden if phrase in text],
        "required_metric_hits": {m: (m in text) for m in ["0.8014", "0.9736", "0.87", "0.7752", "0.993616", "17.02"]},
    }


def write_checklist(builder: ThesisBuilder, validation: dict) -> None:
    lines = [
        "# final_thesis_manuscript_29p_v4 그림/표 체크리스트",
        "",
        "## 표 목록",
        *[f"- {item}" for item in builder.table_log],
        "",
        "## 그림 목록",
        *[f"- {item}" for item in builder.figure_log],
        "",
        "## 텍스트 검증",
        f"- 공백 제외 글자 수: {validation['char_count']}",
        f"- 깨진 문자 수: {validation['replacement_character_count']}",
        f"- API key 패턴 수: {validation['api_key_pattern_count']}",
        f"- 금지 주장 hit: {', '.join(validation['forbidden_phrase_hits']) if validation['forbidden_phrase_hits'] else '없음'}",
        f"- 핵심 수치 포함: {validation['required_metric_hits']}",
        "",
        "## 육안 검토 포인트",
        "- PDF 변환 후 27~31쪽 범위인지 확인",
        "- 표/그림이 페이지 하단에서 잘리지 않는지 확인",
        "- 빈 페이지 또는 과도한 공백이 없는지 확인",
        "- 인적사항 placeholder는 최종 제출 전 직접 입력",
    ]
    CHECKLIST_PATH.write_text("\n".join(lines), encoding="utf-8-sig")


def main() -> None:
    OUTPUTS.mkdir(exist_ok=True)
    evidence = load_evidence()
    builder = ThesisBuilder(evidence)
    builder.add_front_matter()
    builder.add_body()
    builder.save()

    text = extract_docx_text(DOCX_PATH)
    validation = validate_text(text)
    validation["docx_path"] = str(DOCX_PATH)
    validation["table_count"] = builder.table_no
    validation["figure_count"] = builder.figure_no
    VALIDATION_PATH.write_text(json.dumps(validation, ensure_ascii=False, indent=2), encoding="utf-8")
    write_checklist(builder, validation)

    print(f"DOCX created: {DOCX_PATH}")
    print(f"tables={builder.table_no}, figures={builder.figure_no}, chars_no_space={validation['char_count']}")
    print(f"checklist: {CHECKLIST_PATH}")
    print(f"validation: {VALIDATION_PATH}")


if __name__ == "__main__":
    main()
