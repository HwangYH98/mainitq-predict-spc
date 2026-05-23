"""Create a 29-page thesis draft for the capstone project.

The output is intentionally a DOCX draft rather than an HWP file.  The school
cover/approval/abstract HWP templates can be filled by copying from this draft.
"""

from __future__ import annotations

import csv
import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUTS = ROOT / "outputs"
DOCX_PATH = OUTPUTS / "final_thesis_manuscript_29p.docx"
MD_PATH = OUTPUTS / "final_thesis_manuscript_29p.md"

FONT_BODY = "Batang"
FONT_HEAD = "Malgun Gothic"
GREEN = "003D12"
LIGHT_GREEN = "EAF5EE"
LINE = "C9D8CF"


def load_json(path: str) -> dict:
    return json.loads((ROOT / path).read_text(encoding="utf-8"))


def load_csv(path: str) -> list[dict[str, str]]:
    with (ROOT / path).open("r", encoding="utf-8-sig", newline="") as handle:
        return list(csv.DictReader(handle))


def fnum(value, digits: int = 4) -> str:
    try:
        return f"{float(value):.{digits}f}"
    except Exception:
        return str(value)


def pct(value, digits: int = 2) -> str:
    try:
        return f"{float(value) * 100:.{digits}f} %"
    except Exception:
        return str(value)


def pct_compact(value, digits: int = 2) -> str:
    try:
        return f"{float(value) * 100:.{digits}f}%"
    except Exception:
        return str(value)


def collect_evidence() -> dict:
    metrics = load_json("outputs/metrics.json")
    threshold = load_json("outputs/threshold_summary.json")
    spc = load_json("outputs/spc_summary.json")
    ai_ctx = load_json("outputs/ai_report_context.json")
    scania = load_json("outputs/scania_official_cost_metrics.json")
    field = load_json("outputs/field_validation_report.json")
    model_strategy = load_csv("outputs/model_strategy_comparison.csv")
    spc_vs_ml = load_csv("outputs/spc_vs_ml_comparison.csv")
    op_value = load_csv("outputs/operational_value_simulation.csv")
    public_bench = load_csv("outputs/public_industrial_validation_metrics.csv")
    scania_best = next(row for row in scania["metrics"] if row["strategy_id"] == "xgboost_cost_optimized")
    return {
        "metrics": metrics,
        "threshold": threshold,
        "spc": spc,
        "ai_ctx": ai_ctx,
        "scania": scania,
        "scania_best": scania_best,
        "field": field,
        "model_strategy": model_strategy,
        "spc_vs_ml": spc_vs_ml,
        "op_value": op_value,
        "public_bench": public_bench,
    }


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, color: str = "111827", size: int = 9) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(text) < 18 else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(str(text))
    run.bold = bold
    run.font.name = FONT_BODY
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, GREEN)
        set_cell_text(cell, header, bold=True, color="FFFFFF", size=8)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_shading(cells[idx], LIGHT_GREEN if len(table.rows) % 2 else "FFFFFF")
            set_cell_text(cells[idx], value, size=8)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Cm(width)
    doc.add_paragraph()


def add_figure(doc: Document, image_name: str, caption: str, width_in: float = 5.6) -> None:
    path = OUTPUTS / image_name
    if path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(path), width=Inches(width_in))
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cap.runs:
            run.font.size = Pt(9)
            run.font.name = FONT_BODY
            run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = FONT_HEAD
        run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_HEAD)
        run.font.color.rgb = RGBColor.from_string(GREEN)


def add_para(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 2.0
    p.paragraph_format.first_line_indent = Cm(0.7)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = FONT_BODY
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
    run.font.size = Pt(11)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style=None)
        p.paragraph_format.left_indent = Cm(0.7)
        p.paragraph_format.line_spacing = 1.6
        run = p.add_run("· " + item)
        run.font.name = FONT_BODY
        run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
        run.font.size = Pt(10.5)


def page_break(doc: Document) -> None:
    doc.add_page_break()


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(3.5)
    section.left_margin = Cm(3.5)
    section.right_margin = Cm(3.0)
    section.bottom_margin = Cm(2.5)
    section.footer_distance = Cm(1.5)
    style = doc.styles["Normal"]
    style.font.name = FONT_BODY
    style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
    style.font.size = Pt(11)


def thesis_pages(ev: dict) -> list[dict]:
    xgb = ev["metrics"]["models"]["xgboost"]
    logit = ev["metrics"]["models"]["logistic_regression"]
    tuned = ev["threshold"]["selected_metrics"]
    ai_row = ev["ai_ctx"]["row"]
    ai_mode = ev["ai_ctx"].get("report_generation_mode", "-")
    ai_factors = ", ".join(item["feature"].replace("_", " ") for item in ev["ai_ctx"].get("top_shap_factors", [])[:3])
    scania_best = ev["scania_best"]
    balanced_mlspc = next(r for r in ev["op_value"] if r["scenario_id"] == "balanced" and r["policy_id"] == "ml_spc_combined")
    balanced_tuned = next(r for r in ev["op_value"] if r["scenario_id"] == "balanced" and r["policy_id"] == "xgboost_tuned_threshold")

    return [
        {
            "title": "표지 대체 페이지",
            "paras": [
                "ML 예측, Predictive SPC, GenAI 리포트 및 승인형 작업지시를 결합한 스마트 제조 예지보전 운영 시스템 구현",
                "본 문서는 학교 HWP 표지 양식에 옮기기 전 사용할 수 있는 29쪽 내외 논문 본문 초안이다. 표지, 내표지, 인준서, 국문초록은 첨부된 HWP 양식에 맞추어 최종 편집한다.",
                "연구 시스템명: MaintiQ Predict. 데이터 기준: AI4I 2020, SCANIA Component X 공개 benchmark, public run-to-failure adapter. 구현 범위: 로컬 데스크톱 제품형 MVP 및 Streamlit Admin 검증 콘솔.",
            ],
        },
        {
            "title": "국문초록",
            "paras": [
                f"본 연구는 스마트 제조 환경에서 설비 센서 데이터를 기반으로 고장 위험을 예측하고, 예측 결과를 Predictive SPC, GenAI 관리자 리포트, 승인형 작업지시 흐름으로 연결하는 예지보전 운영 시스템을 구현하였다. AI4I 2020 데이터셋을 기준으로 Logistic Regression과 XGBoost를 학습하고, SMOTE 및 threshold tuning을 비교하였다. XGBoost는 PR-AUC {xgb['pr_auc']:.4f}를 보였으며, F1 기준 threshold tuning 결과 임계치는 {ev['threshold']['selected_threshold']:.2f}, precision {tuned['precision']:.4f}, recall {tuned['recall']:.4f}, F1-score {tuned['f1_score']:.4f}로 나타났다.",
                f"또한 ML 위험 확률을 시간축으로 정렬하여 Predictive SPC를 구성하고, Gemini 기반 GenAI 리포트를 통해 관리자 참고용 설명을 생성하였다. 실제 리포트 검증에서는 {ai_mode}로 UDI {ai_row['UDI']} row에 대해 예측 고장 확률 {ai_row['xgboost_probability']:.6f}, 위험 판정 기준 {ai_row['selected_threshold']:.2f}, 상태 {ai_row['risk_status']}를 제시하였다. SCANIA Component X에서는 official class 0~4 cost matrix를 적용하여 rule baseline 대비 official cost metric {float(scania_best['cost_improvement_vs_rule']) * 100:.2f}% 개선 가능성을 확인하였다.",
                "본 연구의 기여는 단일 모델 성능 제시에 그치지 않고 예측, 통계적 모니터링, 설명 리포트, 작업지시 승인, 공개 산업 데이터 검증을 하나의 재현 가능한 로컬 MVP로 통합한 점이다. 단, 실제 PLC/SCADA 운영망 배포, 실제 회사 데이터 성능 재검증, 실제 원화 비용 절감 실증은 본 연구 범위에 포함하지 않는다.",
            ],
        },
        {
            "title": "목차",
            "paras": [
                "1. 서론\n  1.1 연구 배경\n  1.2 연구 목적\n  1.3 연구 범위 및 차별성",
                "2. 이론적 배경 및 선행연구\n  2.1 예지보전과 CBM\n  2.2 SPC와 관리도\n  2.3 ML 기반 고장 예측\n  2.4 설명 가능 AI와 GenAI 리포트\n  2.5 비용 민감 평가",
                "3. 연구 방법\n  3.1 데이터와 전처리\n  3.2 모델 학습과 threshold 정책\n  3.3 Predictive SPC\n  3.4 GenAI 및 작업지시 workflow",
                "4. 시스템 구현\n5. 실험 및 검증\n6. 결론 및 향후 연구\n7. 참고문헌 및 부록",
            ],
        },
        {
            "title": "1. 서론",
            "paras": [
                "스마트 제조 환경에서는 설비의 상태를 센서 데이터로 수집하고, 이를 바탕으로 품질 이상과 설비 고장을 사전에 감지하려는 요구가 증가하고 있다. 기존 제조 현장은 고장 발생 후 수리하는 사후보전 또는 일정 주기마다 점검하는 예방보전에 의존하는 경우가 많다. 그러나 사후보전은 갑작스러운 설비 정지와 생산 손실을 유발할 수 있고, 예방보전은 실제 상태와 무관하게 정비를 수행하므로 과잉 정비와 누락 위험을 동시에 가진다.",
                "예지보전은 설비 센서 데이터를 기반으로 고장 가능성을 예측하여 정비 의사결정을 선제적으로 지원하는 접근이다. 다만 예측 모델이 높은 성능을 보이더라도 그 결과가 현장 담당자가 이해할 수 있는 형태로 제공되지 않거나, 작업지시와 승인 이력으로 연결되지 않으면 실제 운영 가치는 제한된다. 따라서 예측 정확도뿐 아니라 설명성, 모니터링, 의사결정 workflow가 함께 설계되어야 한다.",
            ],
        },
        {
            "title": "1.1 연구 배경",
            "paras": [
                "제조 설비는 온도, 회전 속도, 토크, 공구 마모와 같은 연속형 센서 변수에 의해 상태가 표현된다. 이러한 변수들은 개별적으로도 이상 징후를 가질 수 있지만, 실제 고장 위험은 여러 변수의 조합과 시간적 변화 속에서 나타난다. 따라서 단일 threshold 기반 rule만으로는 복합적인 고장 패턴을 충분히 포착하기 어렵다.",
                "산업공학 관점에서 예지보전은 OEE, MTBF, MTTR과 연결될 수 있다. 고장을 조기에 감지하면 돌발 정지를 줄이고, 정비 준비 시간을 확보하며, 설비 가동률과 정비 계획성을 개선할 수 있다. 본 연구는 이러한 운영 의사결정 관점에서 ML 기반 예측 결과를 Predictive SPC와 작업지시 흐름으로 연결하는 시스템을 구현하였다.",
            ],
        },
        {
            "title": "1.2 연구 목적 및 차별성",
            "paras": [
                "본 연구의 목적은 AI4I 2020 데이터셋을 활용하여 설비 고장 위험을 예측하고, 예측 결과를 관리도 기반 위험 모니터링, GenAI 관리자 리포트, 승인형 작업지시 workflow로 연결하는 제품형 MVP를 구현하는 것이다. 연구 결과는 단순히 모델 성능표로 끝나는 것이 아니라, 사용자가 CSV 데이터를 입력하고 위험 row를 확인하며 리포트와 작업지시를 생성할 수 있는 운영 화면으로 구현된다.",
                "연구의 차별성은 네 가지로 정리된다. 첫째, Logistic Regression과 XGBoost baseline뿐 아니라 SMOTE와 threshold tuning을 비교하였다. 둘째, ML 위험 확률을 시간 흐름으로 보고 Predictive SPC 관점의 관리도와 경보 정책을 구성하였다. 셋째, GenAI 리포트를 통해 예측 결과를 관리자 참고 문장으로 변환하였다. 넷째, SCANIA Component X의 official cost metric을 활용하여 공개 산업 데이터 기준 비용 민감 평가 가능성을 제시하였다.",
            ],
        },
        {
            "title": "2. 이론적 배경 및 선행연구",
            "paras": [
                "예지보전 관련 선행연구는 크게 세 흐름으로 나눌 수 있다. 첫째는 센서 데이터 기반 condition-based maintenance와 predictive maintenance 연구이다. 둘째는 XGBoost, Random Forest, SVM, neural network 등 ML 기반 고장 예측 연구이다. 셋째는 SHAP, LIME 등 설명 가능 AI와 운영 의사결정 지원 도구를 결합하는 연구이다.",
                "Jardine 등은 CBM이 진단, 예후, 정비 의사결정을 포함하는 절차임을 정리하였다. Carvalho 등은 ML 기반 예지보전 연구가 데이터 전처리, 특징 추출, 모델 학습, 성능 평가의 흐름으로 구성된다고 설명하였다. 본 연구는 이러한 흐름을 따르되, 예측 결과를 작업지시 승인 이력까지 연결한다는 점에서 시스템 통합 관점을 강조한다.",
            ],
        },
        {
            "title": "2.1 사후보전, 예방보전, 예지보전",
            "paras": [
                "사후보전은 설비 고장이 발생한 뒤 수리하는 방식이다. 초기 도입 비용은 낮지만, 고장 시점이 예측되지 않아 생산 중단과 긴급 정비 비용이 커질 수 있다. 예방보전은 일정 주기에 따라 점검과 교체를 수행하는 방식으로, 사후보전에 비해 계획성이 높지만 설비 상태를 충분히 반영하지 못하면 불필요한 정비가 발생할 수 있다.",
                "예지보전은 센서 데이터와 고장 이력을 활용하여 설비 상태를 추정하고 고장 위험을 사전에 예측한다. 이 방식은 정비 시점을 실제 상태에 가깝게 조정할 수 있다는 장점이 있으나, 데이터 품질, 모델 설명성, 현장 workflow 연결이 부족하면 실제 적용성이 낮아질 수 있다.",
            ],
            "table": (
                ["구분", "장점", "한계", "본 연구와의 관계"],
                [
                    ["사후보전", "운영 단순", "돌발 정지 위험", "비교 기준"],
                    ["예방보전", "계획 정비 가능", "과잉 정비 가능", "기존 방식"],
                    ["예지보전", "상태 기반 의사결정", "데이터·모델 필요", "본 연구의 핵심"],
                ],
            ),
        },
        {
            "title": "2.2 SPC와 Predictive SPC",
            "paras": [
                "SPC는 공정의 변동을 통계적으로 감시하고 관리 한계를 벗어나는 이상 상태를 탐지하는 방법이다. 일반적으로 중심선, 상한 관리한계(UCL), 하한 관리한계(LCL)를 설정하고 관측값이 관리한계를 벗어나는지 확인한다. 제조 현장에서는 품질 특성치의 안정성을 확인하는 데 널리 사용된다.",
                "본 연구에서 Predictive SPC는 단일 센서값이 아니라 ML 모델이 산출한 고장 위험 확률의 시간 흐름을 감시하는 방식으로 사용된다. AI4I 데이터의 test row를 UDI 순서로 정렬하여 시뮬레이션 시간축을 만들고, 고장 확률과 토크 변수의 관리도 정보를 함께 제시하였다. 이는 실제 실시간 설비 연결은 아니지만, 예측 결과를 운영 모니터링 화면으로 확장하는 구조를 보여준다.",
            ],
            "figure": ("spc_risk_chart.png", "그림 1. ML 위험 확률 기반 Predictive SPC 예시"),
        },
        {
            "title": "2.3 ML 기반 고장 예측",
            "paras": [
                "고장 예측 문제는 대부분 정상 데이터가 많고 고장 데이터가 적은 불균형 이진 분류 문제로 나타난다. 따라서 accuracy만으로 성능을 평가하면 고장 class를 제대로 탐지하지 못해도 높은 값이 나올 수 있다. 본 연구에서는 precision, recall, F1-score, ROC-AUC, PR-AUC를 함께 사용하였다.",
                "Logistic Regression은 해석이 비교적 쉽고 baseline으로 적합하다. XGBoost는 decision tree ensemble 기반의 gradient boosting 방법으로, 탭형 제조 데이터에서 비선형 관계와 변수 간 상호작용을 효과적으로 처리할 수 있다. 본 연구에서는 두 모델을 baseline으로 학습하고, XGBoost를 주요 예측 엔진으로 사용하였다.",
            ],
        },
        {
            "title": "2.4 설명 가능 AI, GenAI 리포트, 비용 민감 평가",
            "paras": [
                "SHAP은 각 feature가 예측 결과에 기여한 정도를 계산하여 black-box 모델의 판단 근거를 설명하는 방법이다. 본 연구에서는 XGBoost 예측 결과에 대해 주요 위험 요인을 추출하고, 이를 GenAI 리포트와 작업지시 초안의 설명 근거로 활용하였다.",
                "비용 민감 평가는 false alarm과 missed failure의 비용 차이를 고려한다. 일반적으로 고장을 놓치는 missed failure는 불필요한 점검보다 큰 비용을 유발할 수 있다. 따라서 본 연구는 threshold tuning뿐 아니라 normalized operating cost simulation과 SCANIA official cost matrix를 활용하여 운영 의사결정 관점의 평가를 수행하였다.",
            ],
        },
        {
            "title": "3. 연구 방법",
            "paras": [
                "본 연구의 전체 절차는 데이터 입력, 전처리, 모델 학습, 예측, 위험 모니터링, GenAI 리포트, 작업지시 승인, 검증 산출물 생성으로 구성된다. 사용자는 제품형 데스크톱 앱에서 센서 CSV를 선택하고, 시스템은 컬럼 매핑과 품질 진단 후 고장 확률과 위험 우선순위를 산출한다.",
                "연구 검증은 별도 Admin 콘솔과 스크립트로 수행하였다. 이 구조는 일반 사용자가 보는 제품 화면과 연구자가 확인하는 실험 근거를 분리하기 위한 것이다. 제품 앱은 운영 흐름을 단순화하고, Admin은 모델 비교, benchmark, field validation template, evidence pack을 확인하는 역할을 가진다.",
            ],
        },
        {
            "title": "3.1 데이터와 전처리",
            "paras": [
                "기본 학습 데이터는 AI4I 2020 데이터셋이다. 총 10,000개의 샘플과 설비 상태 변수, 제품 타입, 고장 여부가 포함되어 있다. target은 Machine failure이며, 본 연구에서는 이진 분류 문제로 정의하였다.",
                "전처리 과정에서는 ID 성격의 UDI와 Product ID를 제거하였다. 또한 TWF, HDF, PWF, OSF, RNF는 세부 고장 유형 또는 target leakage를 유발할 수 있으므로 baseline 학습에서 제외하였다. 범주형 변수 Type은 one-hot encoding으로 변환하였다. train/test split은 stratified 방식으로 수행하여 고장 class 비율을 유지하였다.",
            ],
            "table": (
                ["항목", "처리 방식", "이유"],
                [
                    ["UDI, Product ID", "제거", "식별자 성격으로 일반화에 불필요"],
                    ["Type", "One-hot encoding", "범주형 제품 타입 반영"],
                    ["TWF/HDF/PWF/OSF/RNF", "제거", "고장 유형 leakage 방지"],
                    ["Machine failure", "target", "고장 여부 이진 분류"],
                ],
            ),
        },
        {
            "title": "3.2 모델 학습과 threshold 정책",
            "paras": [
                f"모델은 Logistic Regression과 XGBoost를 기준으로 학습하였다. XGBoost의 default threshold 0.50은 recall을 높이는 반면 false positive가 증가하는 경향이 있었다. 따라서 threshold를 0.05부터 0.95까지 비교하고 F1-score가 가장 높은 threshold를 선택하였다. 최종 선택된 threshold는 {ev['threshold']['selected_threshold']:.2f}이다.",
                "불균형 데이터 대응을 위해 SMOTE 적용 모델도 비교하였다. 다만 실험 결과에서는 XGBoost 기본 모델에 threshold tuning을 적용한 전략이 precision과 recall의 균형 면에서 가장 적합하였다. 이는 oversampling이 항상 성능 향상을 보장하지 않으며, 데이터 특성과 평가 목적에 따라 threshold 정책이 더 직접적인 효과를 가질 수 있음을 보여준다.",
            ],
            "figure": ("threshold_tuning.png", "그림 2. Threshold tuning 결과"),
        },
        {
            "title": "3.3 Predictive SPC와 GenAI 리포트",
            "paras": [
                f"Predictive SPC는 모델 예측 확률을 시간축으로 정렬하여 rolling window 기반으로 위험 흐름을 보여준다. AI4I test prediction {ev['spc']['total_rows']}개 중 threshold 초과 고위험 row는 {ev['spc']['high_risk_count']}개였고, risk control-limit 기준 alert는 {ev['spc']['spc_risk_alert_count']}개였다.",
                f"GenAI 리포트는 위험 context를 관리자 참고 문장으로 변환한다. 실제 검증에서 report_generation_mode는 {ai_mode}였고, UDI {ai_row['UDI']} row에 대해 예측 확률 {ai_row['xgboost_probability']:.6f}, 기준 {ai_row['selected_threshold']:.2f}, 상태 {ai_row['risk_status']}를 요약하였다. 리포트는 현장 확인과 부품 점검을 제안하지만, 자동 정비 명령이 아니라 승인형 작업지시 검토 자료로 제한된다.",
            ],
        },
        {
            "title": "3.4 작업지시 workflow와 field validation 준비",
            "paras": [
                "작업지시 workflow는 센서 이벤트 생성, 작업지시 초안 생성, 작업자 승인 또는 검토 필요 또는 반려 결정, 최근 이력 저장으로 구성된다. 이 구조는 예측 결과를 실제 정비 명령으로 자동 실행하지 않고, 사람이 최종 판단하는 human-in-the-loop 방식이다.",
                "실제 회사 데이터 실증을 위해 labeled sensor CSV, maintenance history CSV, downtime/cost CSV 템플릿을 구성하였다. 세 파일이 모두 제공될 경우 precision, recall, false alarm, missed failure, lead time, downtime delta, maintenance cost delta를 계산할 수 있다. 현재는 실제 회사 로그가 없으므로 field validation report는 claim guardrail을 포함하는 준비 단계로 유지된다.",
            ],
        },
        {
            "title": "4. 시스템 구현",
            "paras": [
                "MaintiQ Predict는 PySide6 기반 Windows 데스크톱 앱으로 구현되었다. 사용자는 브라우저나 Streamlit 서버를 직접 열지 않고 설치형 앱에서 데이터 예측, 위험 모니터링, AI 리포트, 작업지시 기능을 사용할 수 있다. 연구 검증용 Streamlit Admin 콘솔은 별도로 유지하였다.",
                "사용자 앱은 홈, 데이터 예측, 위험 모니터링, AI 리포트, 작업지시 화면으로 구성된다. 화면에는 논문, capstone, stage, PoC와 같은 연구 개발 용어를 노출하지 않고, 실제 운영자가 이해할 수 있는 용어로 정리하였다. API key는 입력 세션에서만 사용하며 파일, README, outputs, Git 기록에 저장하지 않는다.",
            ],
            "figure": ("maintiq_predict_screenshot.png", "그림 3. MaintiQ Predict 데스크톱 앱 화면"),
        },
        {
            "title": "4.1 제품 앱과 Admin 콘솔 분리",
            "paras": [
                "제품 앱은 외부 사용자가 센서 CSV를 넣고 위험도를 확인하는 흐름에 집중한다. 반면 Admin 콘솔은 모델 비교, 공개 benchmark, 산업공학 근거, field validation template, GitHub 업로드 범위 검사 등 연구와 검증에 필요한 정보를 관리한다.",
                "이 분리는 발표와 논문에서 중요하다. 사용자용 제품 화면은 실제 소프트웨어처럼 단순한 업무 흐름을 제공하고, 논문 검증 자료는 Admin과 outputs 문서로 분리되어 재현성을 확보한다. 따라서 제품성과 연구 재현성을 동시에 확보할 수 있다.",
            ],
        },
        {
            "title": "4.2 빠른 점검 모드와 정밀 분석 모드",
            "paras": [
                "설치본은 Full과 Lite로 분리하였다. 빠른 점검 모드는 작은 설치본과 경량 운영 점수를 제공하며, 일반 배포와 시연에 적합하다. 정밀 분석 모드는 XGBoost와 SHAP 기반 분석을 포함하며, 연구 검증과 상세 분석에 사용된다.",
                "두 모드는 계산 방식이 다르므로 결과가 완전히 동일하지 않을 수 있다. 이를 방지하기 위해 결과 CSV에는 engine_profile, score_method, interpretation_note를 기록하였다. 사용자 화면에는 빠른 점검 모드와 정밀 분석 모드의 목적을 짧게 안내하고, 자세한 설명은 README와 Admin 문서에서 제공한다.",
            ],
        },
        {
            "title": "5. 실험 및 검증 설계",
            "paras": [
                "실험은 모델 성능 비교, threshold 정책 비교, SPC-only 대비 ML+SPC 비교, operational cost simulation, GenAI 리포트 생성 검증, SCANIA official cost metric 검증, public benchmark 확장으로 구성하였다. 각 실험은 같은 데이터 split 또는 공식 benchmark split을 기준으로 수행하였다.",
                "주요 평가 지표는 precision, recall, F1-score, ROC-AUC, PR-AUC, alert count, false alarm, missed failure, normalized operating cost이다. SCANIA Component X의 경우 공식 class 0~4 cost matrix를 사용하여 official cost metric을 계산하였다.",
            ],
            "table": (
                ["실험", "비교군", "평가지표"],
                [
                    ["AI4I baseline", "Logistic Regression, XGBoost", "F1, ROC-AUC, PR-AUC"],
                    ["SMOTE/Threshold", "기본, SMOTE, tuned threshold", "precision, recall, F1"],
                    ["SPC 비교", "SPC-only, ML, ML+SPC", "alert, FP, FN"],
                    ["SCANIA", "rule, SPC, multiclass, cost optimized", "official cost"],
                ],
            ),
        },
        {
            "title": "5.1 AI4I baseline 결과",
            "paras": [
                f"AI4I baseline 실험에서 XGBoost는 Logistic Regression보다 높은 PR-AUC를 보였다. Logistic Regression의 PR-AUC는 {logit['pr_auc']:.4f}, XGBoost의 PR-AUC는 {xgb['pr_auc']:.4f}였다. 이는 비선형 관계와 feature interaction을 처리하는 XGBoost가 탭형 제조 데이터에 더 적합할 수 있음을 보여준다.",
                f"XGBoost default threshold 0.50에서는 recall이 {next(r for r in ev['model_strategy'] if r['strategy_id']=='xgboost_default')['recall']}로 높았지만 false positive가 많았다. 반면 F1 기준 tuned threshold 0.87에서는 precision {tuned['precision']:.4f}, recall {tuned['recall']:.4f}, F1-score {tuned['f1_score']:.4f}로 균형이 개선되었다.",
            ],
            "table": (
                ["모델", "Precision", "Recall", "F1", "ROC-AUC", "PR-AUC"],
                [
                    ["Logistic Regression", fnum(logit["precision"]), fnum(logit["recall"]), fnum(logit["f1_score"]), fnum(logit["roc_auc"]), fnum(logit["pr_auc"])],
                    ["XGBoost", fnum(xgb["precision"]), fnum(xgb["recall"]), fnum(xgb["f1_score"]), fnum(xgb["roc_auc"]), fnum(xgb["pr_auc"])],
                ],
            ),
            "figure": ("pr_curve.png", "그림 4. Baseline PR curve"),
        },
        {
            "title": "5.2 SMOTE 및 threshold 비교",
            "paras": [
                "SMOTE는 minority class를 합성하여 불균형 문제를 완화하는 방법이다. 그러나 본 실험에서는 XGBoost에 SMOTE를 적용한 경우 PR-AUC가 0.7163으로 XGBoost 기본 모델의 0.8014보다 낮았다. 이는 SMOTE가 항상 좋은 결과를 보장하지 않으며, 데이터 분포와 모델 특성에 따라 효과가 달라질 수 있음을 보여준다.",
                "threshold tuning은 모델 자체를 바꾸지 않고 decision boundary를 조정하는 방법이다. 본 연구에서는 F1-score 기준으로 threshold 0.87을 선택하였다. 이 정책은 default threshold보다 alert count를 줄이고 precision을 높이면서도 recall을 일정 수준 유지하였다.",
            ],
            "table": (
                ["전략", "Threshold", "Precision", "Recall", "F1", "PR-AUC"],
                [[r["display_name"], r["threshold"], r["precision"], r["recall"], r["f1_score"], r["pr_auc"]] for r in ev["model_strategy"][:4]],
            ),
        },
        {
            "title": "5.3 SPC-only vs ML+SPC 비교",
            "paras": [
                "SPC-only torque control-limit rule은 precision이 높지만 recall이 매우 낮았다. 이는 단일 센서 관리도 기준이 실제 고장 row의 대부분을 놓칠 수 있음을 의미한다. 반면 ML selected threshold는 recall과 precision의 균형이 높았고, ML+Predictive SPC는 recall을 더 높이는 대신 false alarm이 증가하였다.",
                "따라서 운영 정책은 현장의 비용 구조에 따라 선택되어야 한다. 고장을 놓치는 비용이 매우 큰 환경에서는 ML+SPC 결합 정책이 유리할 수 있고, false alarm을 줄이는 것이 중요한 환경에서는 tuned threshold 정책이 더 적합할 수 있다.",
            ],
            "table": (
                ["전략", "Precision", "Recall", "F1", "Alert", "FP", "FN"],
                [[r["display_name"], r["precision"], r["recall"], r["f1_score"], r["alert_count"], r["false_positive"], r["false_negative"]] for r in ev["spc_vs_ml"]],
            ),
            "figure": ("spc_vs_ml_summary.md", ""),
        },
        {
            "title": "5.4 Operational cost simulation",
            "paras": [
                f"운영 cost simulation은 false alarm, missed failure, planned action cost를 상대 단위로 설정하여 정책별 normalized operating cost를 비교하였다. balanced scenario에서 XGBoost tuned threshold의 normalized cost는 {balanced_tuned['normalized_operating_cost']}였고, ML+SPC combined는 {balanced_mlspc['normalized_operating_cost']}였다.",
                "이는 실제 원화 비용 절감 실증이 아니라, 비용 가중치가 주어졌을 때 정책 선택이 운영 비용에 어떤 영향을 줄 수 있는지 평가하는 simulation이다. 따라서 논문에서는 비용 절감 가능성 또는 비용 민감 의사결정 가능성으로 표현해야 하며, 실제 회사 비용 절감으로 표현하면 안 된다.",
            ],
            "figure": ("operational_value_simulation.png", "그림 5. 운영 cost simulation 비교"),
        },
        {
            "title": "5.5 Gemini API 기반 리포트 검증",
            "paras": [
                f"GenAI 리포트 검증에서는 {ai_mode}로 관리자 참고 리포트를 생성하였다. 입력 context는 AI4I UDI {ai_row['UDI']} row였고, 예측 확률은 {ai_row['xgboost_probability']:.6f}, 위험 판정 기준은 {ai_row['selected_threshold']:.2f}, 상태는 {ai_row['risk_status']}였다. 주요 위험 요인은 {ai_factors}로 요약되었다.",
                "리포트는 고위험 상태에 대해 현장 확인, 토크 및 회전속도 관련 부품 점검, 승인형 작업지시 검토를 제안하였다. 이 결과는 예측 모델의 숫자 출력을 관리자에게 이해 가능한 문장으로 변환할 수 있음을 보여준다. 단, 리포트는 AI4I 기반 offline simulation 결과이며 실제 현장 센서 feed 또는 자동 정비 명령이 아니다.",
            ],
            "table": (
                ["항목", "값"],
                [
                    ["report_generation_mode", ai_mode],
                    ["예측 고장 확률", f"{ai_row['xgboost_probability']:.6f}"],
                    ["위험 판정 기준", f"{ai_row['selected_threshold']:.2f}"],
                    ["상태", ai_row["risk_status"]],
                    ["주요 요인", ai_factors],
                ],
            ),
        },
        {
            "title": "5.6 SCANIA official cost metric 검증",
            "paras": [
                f"SCANIA Component X는 실제 SCANIA truck fleet에서 수집된 anonymized multivariate time-series 데이터셋이며, official class 0~4 prediction problem과 cost matrix를 제공한다. 본 연구는 validation data 기준으로 rule baseline, SPC-style baseline, logistic multiclass, XGBoost multiclass, XGBoost cost-optimized 전략을 비교하였다.",
                f"XGBoost cost-optimized 전략은 예측 확률과 official cost matrix를 곱해 expected cost가 가장 낮은 class를 선택한다. 결과적으로 official cost는 {float(scania_best['official_cost']):.0f}, normalized cost는 {float(scania_best['normalized_cost']):.4f}, rule baseline 대비 cost improvement는 {float(scania_best['cost_improvement_vs_rule']) * 100:.2f}%로 나타났다. 이는 실제 원화 비용 절감이 아니라 공개 benchmark official metric 기준 개선 가능성이다.",
            ],
            "figure": ("scania_official_cost_comparison.png", "그림 6. SCANIA official cost metric 비교"),
        },
        {
            "title": "5.7 Public benchmark 확장과 field validation",
            "paras": [
                "본 연구는 SCANIA 외에도 MetroPT-3, NASA C-MAPSS, IMS/FEMTO와 같은 공개 산업 benchmark adapter를 구성하였다. 원본 데이터가 없는 경우 sample smoke로 guardrail을 유지하고, 원본 데이터가 제공되는 경우 lead-time, RUL, early warning, simulated cost 등을 계산할 수 있도록 설계하였다.",
                f"field validation report는 실제 회사 데이터가 없을 때 claim_status를 제한적으로 표시한다. 현재 field_claim_ready 값은 {ev['field'].get('field_claim_ready', False)}이며, 이는 실제 회사 비용 절감 또는 탐지 시간 단축 실증을 주장할 수 없음을 의미한다. 실제 실증을 위해서는 설비 ID, timestamp, 센서값, 고장 라벨, 정비 시작/종료, downtime, 부품비와 인건비 로그가 필요하다.",
            ],
            "table": (
                ["데이터", "검증 가능 항목", "한계"],
                [
                    ["SCANIA", "official cost metric", "실제 회사 원화 비용 아님"],
                    ["MetroPT-3", "이상 이벤트 및 lead-time", "비용 로그 없음"],
                    ["C-MAPSS", "RUL 및 horizon classification", "항공 엔진 benchmark"],
                    ["실제 회사 데이터", "비용·downtime 실증", "현재 미확보"],
                ],
            ),
        },
        {
            "title": "6. 결론",
            "paras": [
                "본 연구는 ML 기반 고장 예측 모델을 구축하고, threshold tuning, Predictive SPC, GenAI 관리자 리포트, 승인형 작업지시 workflow, public benchmark 검증을 하나의 로컬 제품형 MVP로 통합하였다. 실험 결과 XGBoost는 AI4I 기준 PR-AUC 0.8014를 보였고, F1 기준 threshold tuning을 통해 precision과 recall의 균형을 개선하였다.",
                "또한 단일 sensor SPC rule과 ML 기반 경보 정책을 비교함으로써 예측 확률 기반 모니터링의 장점을 확인하였다. Gemini 리포트는 위험 판단을 관리자 참고 문장으로 변환할 수 있음을 보였고, SCANIA official cost metric은 공개 산업 데이터 기준 비용 민감 평가 가능성을 제공하였다.",
            ],
        },
        {
            "title": "6.1 연구 한계 및 향후 연구",
            "paras": [
                "본 연구의 가장 큰 한계는 실제 회사 현장 데이터 실증이 아직 없다는 점이다. 공개 데이터와 simulation을 통해 모델과 workflow의 가능성은 확인했지만, 실제 회사 비용 절감률이나 고장 탐지 시간 단축률을 주장하려면 labeled sensor CSV, 정비 이력, downtime, 비용 로그가 필요하다.",
                "향후 연구에서는 실제 설비 데이터와 정비 로그를 확보하여 before/after 평가를 수행해야 한다. 또한 PLC/SCADA, MQTT, OPC UA 등 실제 현장 연동, 모델 drift monitoring, 운영 DB 전환, 사용자 권한 관리, 코드 서명 및 자동 업데이트 등 상용 배포 품질을 보강할 필요가 있다.",
            ],
        },
        {
            "title": "참고문헌",
            "paras": [
                "[1] Jardine, A. K. S., Lin, D., & Banjevic, D. (2006). A review on machinery diagnostics and prognostics implementing condition-based maintenance. Mechanical Systems and Signal Processing, 20(7), 1483-1510.",
                "[2] Carvalho, T. P., Soares, F. A. A. M. N., Vita, R., Francisco, R. P., Basto, J. P., & Alcalá, S. G. S. (2019). A systematic literature review of machine learning methods applied to predictive maintenance. Computers & Industrial Engineering, 137, 106024.",
                "[3] Montgomery, D. C. (2019). Introduction to Statistical Quality Control. Wiley.",
                "[4] Chen, T., & Guestrin, C. (2016). XGBoost: A scalable tree boosting system. KDD.",
                "[5] Chawla, N. V., Bowyer, K. W., Hall, L. O., & Kegelmeyer, W. P. (2002). SMOTE: Synthetic minority over-sampling technique. Journal of Artificial Intelligence Research, 16, 321-357.",
                "[6] Lundberg, S. M., & Lee, S. I. (2017). A unified approach to interpreting model predictions. NeurIPS.",
                "[7] Elkan, C. (2001). The foundations of cost-sensitive learning. IJCAI.",
                "[8] Matzka, S. (2020). AI4I 2020 Predictive Maintenance Dataset. UCI Machine Learning Repository.",
                "[9] SCANIA Component X Dataset. Researchdata.se.",
                "[10] MetroPT-3 Dataset. UCI Machine Learning Repository.",
                "[11] NASA C-MAPSS and IMS Bearings datasets. NASA Prognostics Center of Excellence.",
                "[12] Google Gemini generateContent API documentation; OpenAI Responses API documentation; PyInstaller and Qt for Python documentation.",
            ],
        },
        {
            "title": "부록 A. 주요 산출물 경로",
            "paras": [
                "본 연구의 재현 가능한 주요 산출물은 outputs 폴더에 저장된다. baseline 결과는 metrics.json, confusion_matrix.png, pr_curve.png, baseline_predictions.csv로 확인할 수 있다. threshold tuning 결과는 threshold_summary.json과 threshold_tuning.png에 저장된다.",
                "GenAI 리포트 검증 근거는 ai_report_context.json, ai_manager_report.md, genai_report_evidence.md에 정리되어 있다. 공개 산업 데이터 검증 결과는 scania_official_cost_metrics.json, public_industrial_validation_metrics.csv, run_to_failure_evidence_summary.md에서 확인할 수 있다. 실제 회사 데이터 실증 준비 자료는 field_data_template.csv, field_maintenance_template.csv, field_cost_template.csv, field_validation_protocol.md에 정리되어 있다.",
            ],
        },
    ]


def build_markdown(pages: list[dict]) -> str:
    lines = ["# MaintiQ Predict 29쪽 내외 논문 초안", ""]
    for idx, page in enumerate(pages, start=1):
        lines.append(f"<!-- page {idx} -->")
        lines.append(f"## {page['title']}")
        lines.extend(page.get("paras", []))
        if "table" in page:
            headers, rows = page["table"]
            lines.append("| " + " | ".join(headers) + " |")
            lines.append("|" + "|".join(["---"] * len(headers)) + "|")
            for row in rows:
                lines.append("| " + " | ".join(row) + " |")
        if "figure" in page and page["figure"][0].endswith(".png"):
            lines.append(f"[그림] {page['figure'][1]}: outputs/{page['figure'][0]}")
        lines.append("")
    return "\n".join(lines)


def build_docx(pages: list[dict]) -> None:
    doc = Document()
    configure_document(doc)

    for idx, page in enumerate(pages, start=1):
        if idx == 1:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run("학사학위논문 초안")
            run.bold = True
            run.font.name = FONT_HEAD
            run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_HEAD)
            run.font.size = Pt(18)
            doc.add_paragraph()
            title_p = doc.add_paragraph()
            title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_run = title_p.add_run(pages[0]["paras"][0])
            title_run.bold = True
            title_run.font.name = FONT_HEAD
            title_run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_HEAD)
            title_run.font.size = Pt(17)
            doc.add_paragraph()
            for para in page["paras"][1:]:
                add_para(doc, para)
        else:
            add_heading(doc, page["title"], 1 if "." not in page["title"] else 2)
            for para in page.get("paras", []):
                if "\n" in para:
                    add_bullets(doc, [part.strip() for part in para.splitlines() if part.strip()])
                else:
                    add_para(doc, para)
            if "table" in page:
                headers, rows = page["table"]
                add_table(doc, headers, rows)
            if "figure" in page:
                image_name, caption = page["figure"]
                if image_name.endswith(".png"):
                    add_figure(doc, image_name, caption)
        footer = doc.sections[-1].footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if not footer.text:
            run = footer.add_run("MaintiQ Predict 논문 초안")
            run.font.size = Pt(9)
            run.font.name = FONT_BODY
            run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
        if idx != len(pages):
            page_break(doc)

    DOCX_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(DOCX_PATH)


def main() -> None:
    ev = collect_evidence()
    pages = thesis_pages(ev)
    # Keep the submitted draft close to the requested 29-page sample format:
    # merge the limitation page into the conclusion and keep references as the
    # final page.  The appendix information remains available in outputs files.
    if len(pages) == 31:
        limitation_page = pages.pop(28)
        pages[27]["paras"].extend(limitation_page["paras"])
        pages.pop(-1)
    if len(pages) != 29:
        raise AssertionError(f"Expected 29 page blocks, found {len(pages)}")
    MD_PATH.write_text(build_markdown(pages), encoding="utf-8")
    build_docx(pages)
    print(f"Markdown: {MD_PATH}")
    print(f"DOCX: {DOCX_PATH}")
    print(f"page_blocks={len(pages)}")


if __name__ == "__main__":
    main()
